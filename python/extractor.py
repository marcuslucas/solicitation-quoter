#!/usr/bin/env python3
"""Text parsing and field extraction for SolicitationQuoter."""

import os, re, json
from pathlib import Path

SCOPE_MAX = 3000


def parse_pdf(filepath):
    try:
        import pdfplumber
        with pdfplumber.open(filepath) as pdf:
            text = "\n\n".join(p.extract_text() or "" for p in pdf.pages).strip()
        if text:
            return text
    except Exception as e:
        print(f"pdfplumber failed: {e}")
    try:
        from pypdf import PdfReader
        reader = PdfReader(filepath)
        text = "\n\n".join(p.extract_text() or "" for p in reader.pages).strip()
        if text:
            return text
    except Exception as e:
        print(f"pypdf failed: {e}")
    return ""


def parse_docx(filepath):
    from docx import Document
    doc = Document(filepath)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def parse_document(filepath):
    """
    Extract text from a document file.

    Delegates to document_loader.load_document() which detects the actual file
    type by magic bytes (not extension). Returns the full text string.
    Raises ValueError if the format is not supported and no text was extracted.
    """
    from document_loader import load_document
    result = load_document(filepath)
    if result.error and not result.text.strip():
        raise ValueError(result.error)
    return result.text


# ── FORMAT DETECTION ──────────────────────────────────────────────────────────

def detect_format(text):
    """
    Identify the solicitation format by scoring multiple fingerprint patterns.

    Formats scored:
      'sam_export'  — SAM.gov structured notice export
      'agency_form' — VA/agency combined synopsis form (SOLICITATION NUMBER*)
      'formal_rfq'  — Formal RFQ with cover page + SECTION A/B/C/D/E
      'sf1449'      — SF-1449 Solicitation/Contract/Order for Commercial Products

    Returns the format name with the highest score when score >= 3, else 'unknown'.
    Scoring beats first-match cascades when multiple formats share keywords
    (e.g. 'SOLICITATION NUMBER' appears in both agency_form and sf1449).
    """
    first_2000 = text[:2000]
    first_5000 = text[:5000]

    scores = {
        "sam_export": 0,
        "agency_form": 0,
        "formal_rfq": 0,
        "sf1449": 0,
    }

    # ── sam_export ────────────────────────────────────────────────────────────
    if "Notice ID:" in first_2000:
        scores["sam_export"] += 3
    if "Combined Synopsis/Solicitation Details" in first_2000:
        scores["sam_export"] += 2
    if "Primary Contact Name:" in first_5000:
        scores["sam_export"] += 1
    if "Notice Details" in first_5000:
        scores["sam_export"] += 1
    # Defensive fingerprints for SAM export variants that use different notice ID labeling
    if "Contracting Office Information" in first_5000:
        scores["sam_export"] += 1
    if re.search(r"Response Date:\s*\d{4}/\d{2}/\d{2}", first_5000):
        scores["sam_export"] += 1
    if "Set Aside Code" in text:
        scores["sam_export"] += 1

    # ── agency_form ───────────────────────────────────────────────────────────
    if re.search(r"SOLICITATION NUMBER\*", first_2000):
        scores["agency_form"] += 3
    if "POINT OF CONTACT*" in text:
        scores["agency_form"] += 2
    if "RESPONSE DATE/TIME/ZONE" in first_2000:
        scores["agency_form"] += 1

    # ── formal_rfq ────────────────────────────────────────────────────────────
    if "Issuing Office:" in first_2000:
        scores["formal_rfq"] += 3
    if re.search(r"\bSECTION\s+[A-E]\b", text):
        scores["formal_rfq"] += 2
    if "Quotation Due Date:" in first_2000:
        scores["formal_rfq"] += 1

    # ── sf1449 ────────────────────────────────────────────────────────────────
    if "STANDARD FORM 1449" in first_5000:
        scores["sf1449"] += 3
    if "Solicitation/Contract/Order for Commercial" in first_2000:
        scores["sf1449"] += 3
    if re.search(r"SECTION\s+I\s+SCHEDULES", text):
        scores["sf1449"] += 2
    if re.search(r"SECTION\s+II\s+CONTRACT\s+CLAUSES", text):
        scores["sf1449"] += 2
    if "SCHEDULE OF SUPPLIES/SERVICES" in first_2000:
        scores["sf1449"] += 1

    best = max(scores, key=scores.get)
    if scores[best] >= 3:
        print(f"[detect_format] scores={scores} -> {best}")
        return best

    print(f"[detect_format] scores={scores} -> unknown")
    return "unknown"


# ── SHARED HELPERS ────────────────────────────────────────────────────────────

def _scope_block(raw):
    """Return scope_of_work dict with optional truncation metadata."""
    raw = re.sub(r"\s+", " ", raw).strip()
    d = {"scope_of_work": raw[:SCOPE_MAX]}
    if len(raw) > SCOPE_MAX:
        d["scope_truncated"] = True
        d["scope_full"] = raw
    else:
        d["scope_truncated"] = False
    return d


# ── FORMAT-SPECIFIC EXTRACTORS ────────────────────────────────────────────────

def extract_sam_export(text):
    """
    Extract fields from SAM.gov-style structured solicitations.
    Fingerprint: 'Notice ID:' near top of page 1.
    Structure: labeled key-value pages with 'Primary Contact', 'Notice Details', etc.
    """
    def find(patterns):
        for p in patterns:
            m = re.search(p, text, re.IGNORECASE | re.MULTILINE)
            if m:
                v = (m.group(1) if m.lastindex else m.group(0)).strip()
                if v: return v
        return ""

    d = {}
    d["solicitation_number"] = find([
        r"Notice\s*ID[:\s]*([A-Z0-9\-]+)",
        r"Solicitation\s*(?:No|Number|#)[:\s]*([A-Z0-9\-]+)",
        r"RFP\s*(?:No|Number|#)[:\s]*([A-Z0-9\-]+)",
        r"RFQ\s*(?:No|Number|#)[:\s]*([A-Z0-9\-]+)",
    ])
    # SAM.gov amendment notices store the base notice ID in "Notice ID:" and the
    # amendment number separately as "Update: NNNN" in the page header.
    # Combine them to reconstruct the full solicitation number (e.g. W911S225U14310001).
    _update_m = re.search(
        r"^[A-Z0-9\-]+:\s+Combined\s+Synopsis/Solicitation\b.*?\bUpdate:\s*(\d+)\b",
        text, re.MULTILINE | re.DOTALL | re.IGNORECASE
    )
    if _update_m and d.get("solicitation_number"):
        suffix = _update_m.group(1)
        if not d["solicitation_number"].endswith(suffix):
            d["solicitation_number"] = d["solicitation_number"] + suffix

    d["project_title"] = find([
        r"Subject[:\s]+(.+?)(?:\n|$)",
        r"(?:Project|Contract|Solicitation)\s+Title[:\s]+(.+?)(?:\n|$)",
    ])
    d["solicitation_type"] = find([
        r"Solicitation\s+Type[:\s]+(\w+)",
        r"\b(RFQ|RFP|IFB|Sources Sought|Combined Synopsis)\b",
    ])
    d["issuing_agency"] = find([
        r"Contracting\s+Office\s+Name[:\s]+(.+?)(?:\n|$)",
        r"(?:Issued by|Issuing Office|Agency)[:\s]+(.+?)(?:\n|$)",
    ])
    d["contracting_office_address"] = find([
        r"Contracting\s+Office\s+Address\s*\n(.+?)(?:\n\n|\Z)",
    ])
    d["due_date"] = find([
        r"Response\s+Date[:\s]*([^\n]+)",
        r"(?:Due Date|Deadline|Response Due)[:\s]*([^\n]+)",
    ])
    d["posting_date"] = find([r"Posting\s+Date[:\s]*([^\n]+)"])
    d["contact_name"]  = find([r"Primary\s+Contact\s+Name[:\s]*([^\n]+)"])
    d["contact_email"] = find([
        r"Primary\s+Contact\s+Email[:\s]*([^\s]+@[^\s]+)",
        r"([a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,})",
    ])
    d["contact_phone"] = find([r"Primary\s+Contact\s+Phone\s*(?:Number)?[:\s]*([^\n]+)"])
    d["naics_code"] = find([r"NAICS\s*(\d{5,6})"])
    d["psc_code"]   = find([r"Product\s+or\s+Service\s+Code\s*([0-9A-Z]+)"])
    d["set_aside"]  = find([
        r"Set\s*Aside\s*Code\s*(.+?)(?:\n|$)",
        r"(Total Small Business Set.?Aside|Small Business|8\(a\)|SDVOSB|HUBZone|WOSB)",
    ])
    d["place_of_performance"] = find([
        r"Place\s+of\s+Performance\s+Address\s*([^\n]+(?:\n[^\n]+)?)",
    ])
    d["period_of_performance"] = find([r"Period\s+of\s+Performance[:\s]+(.+?)(?:\n|$)"])
    d["estimated_value"] = find([
        r"Estimated\s+(?:Value|Cost|Budget)[:\s]*\$?([\d,\.]+(?:\s*(?:million|billion))?)",
        r"Not\s+to\s+Exceed[:\s]*\$?([\d,\.]+)",
    ])

    # Scope of work
    m = re.search(r"Description[:\s]*(.+?)(?=Contact Information|Notice Details|Attachment|\Z)", text, re.IGNORECASE|re.DOTALL)
    raw_scope = m.group(1) if m else text[:SCOPE_MAX * 2]
    d.update(_scope_block(raw_scope))

    # Quantities
    qtys = re.findall(r"\b(SM|S|M|L|XL|XXL|2XL|3XL)[:\s]*(\d+)", text, re.IGNORECASE)
    if qtys:
        d["quantities"] = [{"size": q[0].upper(), "qty": q[1]} for q in qtys]

    # Attachments
    atts = re.findall(r"[•\-\*]?\s*([A-Z0-9_\-]+\.pdf|[A-Z0-9_\-]+\.docx)", text, re.IGNORECASE)
    if atts:
        d["attachments"] = list(set(atts))

    return {k: v for k, v in d.items() if v not in ("", [], None)}


def extract_agency_form(text):
    """
    Extract fields from VA/agency Combined Synopsis/Solicitation form.
    Fingerprint: 'SOLICITATION NUMBER*' on page 1.
    Structure: ALL-CAPS labels with optional * followed by value on the SAME line.
    E.g.: SOLICITATION NUMBER* 36C24225Q0696
    Contact block is multi-line: POINT OF CONTACT* -> blank lines -> role -> name -> email.
    """
    def same_line(label_pat):
        """Match: LABEL[*] VALUE (label and value on same line, space-separated)."""
        m = re.search(label_pat + r'\*?\s+(.+?)(?:\n|$)', text, re.IGNORECASE)
        return m.group(1).strip() if m else ""

    d = {}
    d["solicitation_number"] = same_line(r"SOLICITATION NUMBER")
    d["project_title"]       = same_line(r"SUBJECT")
    # Strip trailing city/country suffix (e.g. ", NEW YORK, USA") — stop at first comma
    due_raw = same_line(r"RESPONSE DATE/TIME/ZONE")
    d["due_date"] = re.sub(r",.*$", "", due_raw).strip()
    d["issuing_agency"]      = same_line(r"CONTRACTING OFFICE ADDRESS")

    naics_m = re.search(r"NAICS CODE\*?\s+(\d{5,6})", text, re.IGNORECASE)
    if naics_m:
        d["naics_code"] = naics_m.group(1).strip()

    psc_m = re.search(r"PRODUCT SERVICE CODE\*?\s+([A-Z0-9]{3,4})", text, re.IGNORECASE)
    if psc_m:
        d["psc_code"] = psc_m.group(1).strip()

    # Contact name — find the line immediately before the email address in the
    # POINT OF CONTACT block. Works for both layouts:
    #   pdfplumber: "POINT OF CONTACT* Contract Officer\nNathan Northrup\nemail"
    #   pypdf:      "POINT OF CONTACT*\n\n\nContract Officer\nNathan Northrup\nemail"
    poc_m = re.search(r"POINT OF CONTACT\*?[^\n]*\n(.+?)(?=PLACE OF PERFORMANCE|\Z)", text, re.DOTALL)
    if poc_m:
        name_m = re.search(
            r"([A-Z][a-z]+(?:\s+[A-Z][a-z.]+)+)\s*\n\s*([\w.%+\-]+@[\w.\-]+\.\w{2,})",
            poc_m.group(1))
        if name_m:
            d["contact_name"] = name_m.group(1).strip()

    # Email — prefer dedicated label (apostrophe may be a replacement char in pypdf)
    email_m = re.search(r"AGENCY CONTACT.{1,3}S EMAIL ADDRESS\s+([\w.%+\-]+@[\w.\-]+\.\w{2,})", text, re.IGNORECASE)
    if not email_m:
        email_m = re.search(r"([\w.%+\-]+@[\w.\-]+\.\w{2,})", text)
    if email_m:
        d["contact_email"] = email_m.group(1).strip()

    # Place of performance — full address block until next ALL-CAPS label
    pop_m = re.search(
        r"PLACE OF PERFORMANCE\s*\nADDRESS\s+(.+?)(?=\nPOSTAL CODE|\nCOUNTRY|\nADDITIONAL|\Z)",
        text, re.DOTALL)
    if pop_m:
        zip_m = re.search(r"POSTAL CODE\s+(\d{5})", text)
        lines = [re.sub(r"\s+and\s*$", "", l.strip(), flags=re.IGNORECASE)
                 for l in pop_m.group(1).split('\n') if l.strip()]
        val = ", ".join(lines)
        if zip_m:
            val = val + " " + zip_m.group(1)
        d["place_of_performance"] = val

    # Solicitation type from prose
    type_m = re.search(r"solicitation is issued as (?:an?\s+)?([A-Z]+)", text, re.IGNORECASE)
    if type_m:
        d["solicitation_type"] = type_m.group(1).strip()
    else:
        type_m2 = re.search(r"\b(RFQ|RFP|IFB)\b", text)
        if type_m2:
            d["solicitation_type"] = type_m2.group(1)

    # Scope — description paragraph from prose
    scope_m = re.search(
        r"(?:is seeking|firm fixed.price service contract for)\s+(.+?)"
        r"(?:\n\n|All interested|See attached|\Z)",
        text, re.IGNORECASE | re.DOTALL)
    if scope_m:
        d.update(_scope_block(scope_m.group(1)))
    else:
        d.update(_scope_block(text[:SCOPE_MAX * 2]))

    return {k: v for k, v in d.items() if v not in ("", [], None)}


def extract_formal_rfq(text):
    """
    Extract fields from formal RFQ documents with cover page + lettered sections.
    Fingerprint: 'Issuing Office:' on page 1, or SECTION A/B/C/D/E headings.
    Structure: labeled cover page fields; NAICS/PSC in Section A prose; SOW in Section C.
    """
    def labeled(label_pat):
        """Match: Label: Value (same line)."""
        m = re.search(label_pat + r'\s*(.+?)(?:\n|$)', text, re.IGNORECASE)
        return m.group(1).strip() if m else ""

    d = {}
    d["solicitation_number"] = labeled(r"Solicitation Number:")
    d["project_title"]       = labeled(r"Title:")
    d["posting_date"]        = labeled(r"Solicitation Release Date:")
    d["solicitation_type"]   = "RFQ"

    # Due date — strip leading weekday name if present
    due_m = re.search(
        r"Quotation Due Date:\s*(?:(?:Monday|Tuesday|Wednesday|Thursday|Friday|Saturday|Sunday),\s*)?(.+?)(?:\n|$)",
        text, re.IGNORECASE)
    if due_m:
        d["due_date"] = due_m.group(1).strip().rstrip(',')

    # Issuing agency — multi-line block; stop before street address / ZIP
    agency_m = re.search(r"Issuing Office:\s+(.+?)(?:\n\n|\Z)", text, re.DOTALL)
    if agency_m:
        lines = [l.strip() for l in agency_m.group(1).split('\n') if l.strip()]
        org_lines = []
        for line in lines:
            if re.match(r'\d+\s+\w', line) or re.search(r'\b[A-Z]{2}\s+\d{5}', line):
                break
            org_lines.append(line)
        if org_lines:
            d["issuing_agency"] = ", ".join(org_lines)

    # Agency contact — "Agency Contact: Name, Title, (XXX) XXX-XXXX, email"
    contact_m = re.search(r"Agency Contact:\s+([^,\n]+)", text, re.IGNORECASE)
    if contact_m:
        d["contact_name"] = contact_m.group(1).strip()

    phone_m = re.search(r"Agency Contact:[^\n]*\((\d{3})\)\s*(\d{3})-(\d{4})", text, re.IGNORECASE)
    if phone_m:
        d["contact_phone"] = phone_m.group(1) + phone_m.group(2) + phone_m.group(3)

    email_m = re.search(r"Agency Contact:[^\n]*([\w.%+\-]+@[\w.\-]+\.\w{2,})", text, re.IGNORECASE)
    if email_m:
        d["contact_email"] = email_m.group(1).strip()

    # NAICS / PSC from Section A prose
    naics_m = re.search(r"NAICS[^\d]{0,30}(\d{5,6})", text, re.IGNORECASE)
    if naics_m:
        d["naics_code"] = naics_m.group(1).strip()

    psc_m = re.search(r"Product Service Code \(PSC\)\s+(?:is\s+)?([A-Z]\d{3,4})", text, re.IGNORECASE)
    if psc_m:
        d["psc_code"] = psc_m.group(1).strip()

    set_aside_m = re.search(r"TOTAL SMALL BUSINESS SET.?\s*ASIDE", text, re.IGNORECASE)
    if set_aside_m:
        d["set_aside"] = "Total Small Business Set-Aside"

    # Place of performance — "located at ADDRESS" in SOW prose.
    # DOTALL needed because pypdf may split "5th" across lines (superscript artifact).
    # Rejoin lowercase continuations after capture (e.g. "5\nth" → "5th").
    pop_m = re.search(r"located at\s+(.+?)(?:\.|,\s*and\b|\Z)", text, re.IGNORECASE | re.DOTALL)
    if pop_m:
        val = re.sub(r"\n([a-z])", r"\1", pop_m.group(1))
        d["place_of_performance"] = re.sub(r"\s+", " ", val).strip()

    # Period of performance — Base Period + all Option lines that follow
    base_m = re.search(r"Base Period:\s*(.+?)(?:\n|$)", text, re.IGNORECASE)
    if base_m:
        base_val = base_m.group(1).strip()
        opts = []
        for opt_line in re.findall(r"Option\s+\d+:[^\n]+", text, re.IGNORECASE):
            dates = re.findall(r"\d{2}/\d{2}/\d{4}", opt_line)
            if dates:
                opts.append(dates[-1])  # end date is the last date on the line
        if opts:
            n = len(opts)
            d["period_of_performance"] = (
                f"Base period {base_val}, plus {n} option period"
                + ("s" if n != 1 else "")
                + f" of 12 months each through {opts[-1]}"
            )
        else:
            d["period_of_performance"] = "Base period " + base_val

    # Scope — full Section C content
    scope_m = re.search(
        r"SECTION C[^\n]*\n(.+?)(?=SECTION\s+[D-Z]\b|\Z)",
        text, re.IGNORECASE | re.DOTALL)
    if scope_m:
        d.update(_scope_block(scope_m.group(1)))
    else:
        d.update(_scope_block(text[:SCOPE_MAX * 2]))

    return {k: v for k, v in d.items() if v not in ("", [], None)}


def extract_sf1449(text):
    """
    Extract fields from SF-1449 Solicitation/Contract/Order for Commercial Products.

    Fingerprint: 'STANDARD FORM 1449' or 'Solicitation/Contract/Order for Commercial'
    Structure: Numbered blocks on page 1 (5=sol#, 6=issue date, 7=contact/due date,
    9=issuing office, 10=set-aside/NAICS, 14=method), then SECTION I/II/III.

    All patterns use re.IGNORECASE | re.MULTILINE.
    Returns a dict; keys with empty/None values are omitted.
    """
    def find(patterns):
        for p in patterns:
            m = re.search(p, text, re.IGNORECASE | re.MULTILINE | re.DOTALL)
            if m:
                v = (m.group(1) if m.lastindex else m.group(0)).strip()
                if v:
                    return v
        return ""

    d = {}

    # Block 5 — Solicitation Number
    # Layout: "5.SOLICITATION NUMBER 6.SOLICITATION\nISSUE DATE\n<SOL_NUM> <DATE>"
    d["solicitation_number"] = find([
        r"5\.SOLICITATION\s+NUMBER\s+6\.SOLICITATION\s*\n\s*ISSUE\s+DATE\s*\n\s*([A-Z0-9]{10,})\s",
        r"(?:SOLICITATION\s+NUMBER)\s*\n?\s*([A-Z0-9]{10,})\b",
        # Fallback: any alphanumeric token that looks like a solicitation number near page start
        r"\b([A-Z]{1,6}\d{2,}[A-Z0-9\-]{2,})\b",
    ])

    # Block 6 — Solicitation Issue Date (posting date)
    # Layout: same line as solicitation number: "<SOL_NUM> <MM/DD/YYYY>"
    d["posting_date"] = find([
        r"ISSUE\s+DATE\s*\n\s*\S+\s+(\d{2}/\d{2}/\d{4})",
        r"SOLICITATION\s+ISSUE\s+DATE\s*\n?\s*\S+\s+(\d{2}/\d{2}/\d{4})",
    ])

    # Block 7 — Contact name and email
    # Layout: "INFORMATION CALL: Crockett, John john.t.crockett@cbp.dhs.gov 05/15/2026 2:00PM ET"
    contact_m = re.search(
        r"INFORMATION\s+CALL:\s*(.*?)\s+([\w.%+\-]+@[\w.\-]+\.\w{2,})",
        text, re.IGNORECASE
    )
    if contact_m:
        d["contact_name"] = contact_m.group(1).strip()
        d["contact_email"] = contact_m.group(2).strip()

    if not d.get("contact_email"):
        d["contact_email"] = find([r"([\w.%+\-]+@[\w.\-]+\.(?:gov|mil|us|com))\b"])

    # Block 8 — Offer Due Date / Local Time
    # Layout A (70B): "... <email> MM/DD/YYYY H:MMxM TZ" on same line as contact
    # Layout B (18Q0042): "OFFER DUE DATE/LOCAL TIME\nINFO CALL: ... HH:MM AM DD Mon YYYY"
    d["due_date"] = find([
        r"(\d{2}/\d{2}/\d{4}\s+\d{1,2}:\d{2}[AP]M\s+\w+)",
        r"OFFER\s+DUE\s+DATE[^\n]*\n.*?(\d{2}/\d{2}/\d{4})",
        r"OFFER\s+DUE\s+DATE[^\n]*\n[^\n]*?(\d{1,2}:\d{2}\s+[AP]M\s+\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4})",
    ])

    # Block 9 — Issued By (issuing agency)
    # Layout (pdfplumber): "ISSUED BY CODE 7014 10.THIS ACQUISITION...\nDHS - Customs & Border Protection SMALL BUSINESS...\n...\nMission Support Contracting Division NAICS:"
    # Strategy: extract org name line directly, then find division line.
    agency_m = re.search(
        r"\n([A-Z][A-Z\s\-&]+(?:Protection|Agency|Command|Service|Office|Bureau))\s+(?:SMALL|WOMEN|HUBZONE|ELIGIBLE|UNRESTRICTED)",
        text, re.IGNORECASE
    )
    if agency_m:
        line1 = agency_m.group(1).strip()
        # Find the contracting division line (mixed case, before "NAICS:")
        div_m = re.search(
            r"([A-Z][a-z].+?(?:Division|Office|Command|Department|Directorate))\s+(?:NAICS|CODE)",
            text, re.IGNORECASE
        )
        if div_m:
            d["issuing_agency"] = line1 + ", " + div_m.group(1).strip()
        else:
            d["issuing_agency"] = line1

    # Block 10 — Set-Aside percentage and NAICS
    # Layout: "SET ASIDE : 100 % FOR:\nDHS - ... SMALL BUSINESS WOMEN-OWNED..."
    sa_m = re.search(r"SET\s+ASIDE\s*:\s*(\d+)\s*%\s*FOR", text, re.IGNORECASE)
    if sa_m:
        d["set_aside"] = f"Small Business Set-Aside {sa_m.group(1)}%"

    # NAICS — appears after "NAICS:\n" with noise lines in between
    d["naics_code"] = find([
        r"NAICS:\s*\n.*?(\d{6})",
        r"NAICS[:\s]+(\d{5,6})\b",
    ])

    # Block 14 — Method of solicitation (RFQ / IFB / RFP)
    # Prefer prose reference ("request for proposal (RFP)") over the checkbox list
    d["solicitation_type"] = find([
        r"request\s+for\s+proposal\s+\(([A-Z]+)\)",
        r"request\s+for\s+quotation\s+\(([A-Z]+)\)",
        r"invitation\s+for\s+bid\s+\(([A-Z]+)\)",
        r"METHOD\s+OF\s+SOLICITATION\s*\n?\s*(RFQ|IFB|RFP)\b",
    ])

    # Project title — item description line in SCHEDULE OF SUPPLIES/SERVICES table
    # Layout: "ITEM NUMBER SCHEDULE OF SUPPLIES/SERVICES QUANTITY UNIT...\n10 <Title>"
    d["project_title"] = find([
        r"SCHEDULE\s+OF\s+SUPPLIES/SERVICES\s+QUANTITY[^\n]*\n\s*\d+\s+(.+?)(?:\n|$)",
        r"(?:Subject|Title)[:\s]+(.+?)(?:\n|$)",
    ])

    # ── Section I content (actual body, not TOC reference) ───────────────────
    # TOC lines have dotted fillers; actual section starts "SECTION I SCHEDULES\n"
    sec_i_m = re.search(r"SECTION I SCHEDULES\s*\n(.+?)(?=SECTION II)", text, re.DOTALL | re.IGNORECASE)
    sec_i = sec_i_m.group(1) if sec_i_m else ""

    # I.2 Minimum Guarantee
    mg_m = re.search(r"MINIMUM\s+GUARANTEE[^$]*\$([\d,]+\.?\d*)", sec_i or text, re.IGNORECASE)
    if mg_m:
        d["minimum_guarantee"] = "$" + mg_m.group(1)

    # I.3 Maximum Amount (estimated value)
    # Pattern: "not exceed a total of $X"
    ev_m = re.search(
        r"(?:not\s+exceed[^$]*|MAXIMUM\s+AMOUNT[^\n]*\n.*?not\s+exceed[^$]*)\$([\d,]+\.?\d*)",
        sec_i or text, re.IGNORECASE | re.DOTALL
    )
    if ev_m:
        d["estimated_value"] = "$" + ev_m.group(1)

    # I.5 Period of Performance
    pop_m = re.search(r"I\.5\s+PERIOD\s+OF\s+PERFORMANCE[:\s]+(.+?)(?=I\.6)", sec_i, re.DOTALL | re.IGNORECASE)
    if pop_m:
        raw_pop = re.sub(r"\s+", " ", pop_m.group(1)).strip()
        d.update(_scope_block.__func__(raw_pop) if hasattr(_scope_block, '__func__') else {})
        d["period_of_performance"] = raw_pop[:500]  # cap at 500 chars

    # Contract type: IDIQ + FFP
    has_idiq = bool(re.search(r"\b(?:IDIQ|ID/IQ|indefinite\s+delivery)", text, re.IGNORECASE))
    has_ffp = bool(re.search(r"\bfirm\s+fixed\s+price\b", text, re.IGNORECASE))
    if has_idiq and has_ffp:
        d["contract_type"] = "IDIQ with Firm Fixed Price delivery orders"
    elif has_idiq:
        d["contract_type"] = "IDIQ"

    # Scope of work — ADDITIONAL INFORMATION (actual requirement) preferred over I.1 (FAR boilerplate)
    scope_src = sec_i if sec_i else text
    scope_m = re.search(
        r"ADDITIONAL\s+INFORMATION[:\s]*\n(.+?)(?=\n32[a-z]\.|Page\s+\d|\Z)",
        text, re.DOTALL | re.IGNORECASE
    )
    if not scope_m:
        scope_m = re.search(
            r"I\.1\s+DESCRIPTION[:\s]+(.+?)(?=I\.2|MINIMUM\s+GUARANTEE|\Z)",
            scope_src, re.DOTALL | re.IGNORECASE
        )
    raw_scope = scope_m.group(1) if scope_m else scope_src[:SCOPE_MAX * 2]
    d.update(_scope_block(raw_scope))

    return {k: v for k, v in d.items() if v not in ("", [], None)}


# ── GENERIC FALLBACK ──────────────────────────────────────────────────────────

def apply_generic_fallback(d, text):
    """Fill any still-empty fields using format-agnostic patterns. Modifies d in-place."""
    def try_fill(key, patterns):
        if d.get(key):
            return
        for p in patterns:
            m = re.search(p, text, re.IGNORECASE | re.MULTILINE)
            if m:
                v = (m.group(1) if m.lastindex else m.group(0)).strip()
                if v:
                    d[key] = v
                    return

    try_fill("solicitation_number", [
        r"(?:solicitation\s*(?:number|#|no\.?))\s*[:\s]*([A-Z0-9][-A-Z0-9]{6,})",
    ])
    try_fill("naics_code", [
        r"NAICS[^\d]{0,20}(\d{5,6})",
    ])
    try_fill("psc_code", [
        r"(?:PSC|Product\s+Service\s+Code)[^\w]{0,20}([A-Z]?\d{3,4}[A-Z]?)",
    ])
    try_fill("contact_email", [
        r"([\w.%+\-]+@[\w.\-]+\.\w{2,})",
    ])
    try_fill("due_date", [
        r"(?:response|due|deadline|closing)[^0-9]{0,30}(\d{1,2}[-/]\d{1,2}[-/]\d{2,4}[^,\n]*)",
    ])
    try_fill("set_aside", [
        r"(Total Small Business Set.?Aside)",
        r"(Small Business Set.?Aside)",
        r"(8\(a\)(?:\s+set.?aside)?)",
        r"(SDVOSB)",
        r"(HUBZone)",
        r"(WOSB)",
    ])


# ── AI EXTRACTION ─────────────────────────────────────────────────────────────

def ai_extract(text):
    api_key = os.environ.get('ANTHROPIC_API_KEY', '')
    if not api_key:
        raise ValueError("No API key configured — set ANTHROPIC_API_KEY environment variable")
    import anthropic
    client = anthropic.Anthropic(api_key=api_key)
    if len(text) > 14000:
        text = text[:7000] + "\n\n[...truncated...]\n\n" + text[-7000:]
    prompt = f"""Extract all data from this solicitation and return ONLY a JSON object with these keys (empty string if not found):
solicitation_number, project_title, solicitation_type, issuing_agency, contracting_office_address,
due_date, posting_date, contact_name, contact_email, contact_phone, naics_code, psc_code, set_aside,
place_of_performance, contract_type, period_of_performance, estimated_value, scope_of_work,
quantities (array of {{size,qty}}), special_requirements (array of strings), attachments (array of strings),
line_items (array of {{description, size, unit, qty, unit_price}} — extract from CLIN tables or item lists if present; use empty string for unknown unit_price)

SOLICITATION:
{text}"""
    resp = client.messages.create(model="claude-sonnet-4-6", max_tokens=2000,
                                   messages=[{"role":"user","content":prompt}])
    raw = re.sub(r"^```(?:json)?\s*","",resp.content[0].text.strip())
    raw = re.sub(r"\s*```$","",raw)
    return json.loads(raw)


# ── ENTRY POINT ───────────────────────────────────────────────────────────────

def extract_data(text, api_key=""):
    format_name = detect_format(text)
    print(f"Detected format: {format_name}")

    if format_name == "sam_export":
        d = extract_sam_export(text)
    elif format_name == "agency_form":
        d = extract_agency_form(text)
    elif format_name == "formal_rfq":
        d = extract_formal_rfq(text)
    elif format_name == "sf1449":
        d = extract_sf1449(text)
    else:
        d = {}

    apply_generic_fallback(d, text)
    d["_format"] = format_name
    d["_method"] = "rules"

    _env_key = os.environ.get('ANTHROPIC_API_KEY', '')
    if api_key or _env_key:
        try:
            ai = ai_extract(text)
            merged = {**d}
            for k, v in ai.items():
                if v and v != "" and v != [] and v != {}:
                    merged[k] = v
            if ai.get("line_items"):
                merged["ai_line_items"] = ai["line_items"]
            merged["_method"] = "ai+rules"
            return merged
        except Exception as e:
            print(f"AI failed, using rules: {e}")

    return d


# Keep the old name as an alias so any direct callers don't break
def extract(text):
    return extract_sam_export(text)


# ── SOW / XLSX LINE ITEM EXTRACTION ──────────────────────────────────────────

# Matches terminal SOW section headers: "4.X.Y. Title." or "4.X.Y Title."
# Trailing period after section number is optional (some sections omit it).
# Title-ending period is optional; lookahead for newline handles titles ending with ')'.
_SOW_SECTION_RE = re.compile(
    r"^(4\.\d{1,2}\.\d{1,2})\.?\s+"
    r"((?:[^.()\n]|\([^)]*\))+)"
    r"(?:\.|(?=[\r\n]))",          # <-- was: (?:\.(?=\n))
    re.MULTILINE
)

# Manufacturer reference patterns.
# Use \s+ between words to tolerate line-wraps inside the phrase.
_PAT_DEF_TECH = re.compile(
    r"equal\s+to\s+or\s+(?:better|greater)\s+than\s+([\w\s]+?)\s+part\s+number\s+([\w\-/]+)",
    re.IGNORECASE
)
# Avon mask model: "equal to or better than the Avon PC50, Protective Mask"
_PAT_AVON_MODEL = re.compile(
    r"equal\s+to\s+or\s+(?:better|greater)\s+than\s+the\s+Avon\s+([\w]+),\s+Protective",
    re.IGNORECASE
)
# Avon part-number: "...Avon Clear Outsert, #70501-156" or "...Avon CTCF50..., 72606/3"
_PAT_AVON_PART = re.compile(
    r"equal\s+to\s+or\s+(?:better|greater)\s+than\s+the\s+Avon\s+\w+.*?[,#]\s*([\w\-/]+)",
    re.IGNORECASE | re.DOTALL
)
# Fallback: "part number X" anywhere in body (for unusual phrasing like 4.6.8)
_PAT_PART_NUM_BARE = re.compile(r"part\s+number\s+([\w\-/]+(?:\s+[A-Z]{1,4})?)", re.IGNORECASE)


def _normalize_sow_title(title):
    """Clean pdfplumber encoding artifacts from SOW section titles.

    The PDF uses smart-quote chars (U+201C/U+201D) as bracket substitutes for
    some parenthesised phrases, and U+2013 en-dashes for range separators.
    """
    # Paired smart-quotes → parens: "Text" → (Text)
    title = re.sub('\u201c([^\u201c\u201d.()\n]+)\u201d', r'(\1)', title)
    # Lone opening smart-quote → open paren; auto-balance
    if '\u201c' in title:
        title = title.replace('\u201c', '(')
        diff = title.count('(') - title.count(')')
        if diff > 0:
            title += ')' * diff
    # Lone closing smart-quote → close paren
    title = title.replace('\u201d', ')')
    # U+FFFD replacement char flanked by spaces → en-dash
    title = re.sub('\ufffd', '\u2013', title)
    # Strip leading "The " article (e.g. "The Low Roll Reloadable Steel Body")
    title = re.sub(r'^The\s+', '', title)
    return title.strip()


def extract_sow_line_items(text, page_texts=None):
    """
    Extract terminal line items from a DHS LLSM Statement of Work PDF text.

    Parses 4.X.Y section headers; handles optional trailing periods and
    pdfplumber encoding artifacts. Extracts description, manufacturer_ref,
    part_number, unit (EA/QT/GA/KT). qty and unit_price are placeholders.

    Returns list of dicts sorted by sow_section.
    """
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    matches = list(_SOW_SECTION_RE.finditer(text))

    page_map = []
    if page_texts:
        pos = 0
        for i, pt in enumerate(page_texts):
            normalized_pt = pt.replace('\r\n', '\n').replace('\r', '\n')
            page_map.append((pos, pos + len(normalized_pt), i + 1))
            pos += len(normalized_pt) + 2  # matches the \n\n join in document_loader

    def find_page(char_pos):
        for start, end, page_num in page_map:
            if start <= char_pos < end:
                return page_num
        return None
    if not matches:
        return []

    # Count raw title occurrences for disambiguation (two-pass: count first)
    title_counts = {}
    for m in matches:
        raw = _normalize_sow_title(m.group(2).strip())
        title_counts[raw] = title_counts.get(raw, 0) + 1

    items = []
    for i, m in enumerate(matches):
        section_num = m.group(1)
        title = _normalize_sow_title(m.group(2).strip())

        # True body: from end of section header to start of next section header
        body_start = m.end()
        body_end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        body = text[body_start:body_end]
        body_lower = body.lower()

        # Unit from body text
        if "1-quart" in body_lower:
            unit = "QT"
        elif "1-gallon" in body_lower:
            unit = "GA"
        elif section_num.startswith("4.8."):
            unit = "KT"
        else:
            unit = "EA"

        # Manufacturer reference and part number
        manufacturer_ref = None
        part_number = None

        m_dt = _PAT_DEF_TECH.search(body)
        if m_dt:
            manufacturer_ref = "Defense Technologies"
            raw_pn = m_dt.group(2).strip()
            # Handle PDF space in part numbers: "8922 NR" → "8922NR"
            next_text = body[m_dt.end(2):]
            sfx = re.match(r'\s+([A-Z]{1,4})\b', next_text)
            if sfx and re.match(r'^[A-Z]+$', sfx.group(1)):
                raw_pn = raw_pn + sfx.group(1)
            part_number = raw_pn
        else:
            m_am = _PAT_AVON_MODEL.search(body)
            if m_am:
                manufacturer_ref = "Avon"
                part_number = m_am.group(1).strip()
            else:
                m_ap = _PAT_AVON_PART.search(body)
                if m_ap:
                    manufacturer_ref = "Avon"
                    part_number = m_ap.group(1).strip().lstrip('#')
                elif "defense technologies" in body_lower:
                    # Fallback for unusual phrasing (e.g. "reload, part number 7001CI")
                    m_pn = _PAT_PART_NUM_BARE.search(body)
                    if m_pn:
                        manufacturer_ref = "Defense Technologies"
                        part_number = re.sub(r'\s+', '', m_pn.group(1).strip())

        # Disambiguation: titles appearing more than once get a body-derived suffix
        if title_counts.get(title, 0) > 1:
            if title == "Liquid Smoke Solution":
                if "1-quart" in body_lower:
                    title = title + " (1 Quart)"
                elif "1-gallon" in body_lower:
                    title = title + " (1 Gallon)"
            elif title == "Respiratory Protection Mask" and part_number:
                title = title + f" ({part_number})"
            elif title == "40mm Multiple Rubber Ball Round":
                if "60 caliber" in body_lower:
                    title = title + " (60 cal)"
                elif "32 caliber" in body_lower:
                    title = title + " (32 cal)"
            elif title == "40mm Multiple Rubber Ball Round (Shortened)":
                if "60 caliber" in body_lower:
                    title = "40mm Multiple Rubber Ball Round (60 cal) (Shortened)"
                elif "32 caliber" in body_lower:
                    title = "40mm Multiple Rubber Ball Round (32 cal) (Shortened)"

        items.append({
            "sow_section": section_num,
            "description": title,
            "manufacturer_ref": manufacturer_ref,
            "part_number": part_number,
            "unit": unit,
            "qty": "N/A",
            "unit_price": "N/A",
            "spec_text": body.strip()[:2000],
            "source_page": find_page(m.start()) if page_map else None,
        })

    def _section_key(item):
        return tuple(int(p) for p in item["sow_section"].split(".") if p.isdigit())

    items.sort(key=_section_key)
    return items


def _iter_xlsx_rows(filepath):
    """Yield rows as lists of cell values. Uses openpyxl if installed; stdlib otherwise."""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(filepath, read_only=True, data_only=True)
        ws = wb.active
        for row in ws.rows:
            yield [cell.value for cell in row]
        wb.close()
        return
    except ImportError:
        pass

    # stdlib fallback: zipfile + xml.etree.ElementTree
    import zipfile
    import xml.etree.ElementTree as ET
    _NS = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"

    with zipfile.ZipFile(filepath) as zf:
        # Shared strings
        shared_strings = []
        if "xl/sharedStrings.xml" in zf.namelist():
            ss_root = ET.parse(zf.open("xl/sharedStrings.xml")).getroot()
            for si in ss_root:
                parts = [elem.text for elem in si.iter(f"{{{_NS}}}t") if elem.text]
                shared_strings.append("".join(parts))

        # Worksheet
        ws_root = ET.parse(zf.open("xl/worksheets/sheet1.xml")).getroot()
        for row_elem in ws_root.iter(f"{{{_NS}}}row"):
            cells = {}
            for c in row_elem:
                ref = c.get("r", "")
                col_str = re.sub(r"\d", "", ref).upper()
                col_idx = sum((ord(ch) - ord('A') + 1) * (26 ** (len(col_str) - 1 - k))
                              for k, ch in enumerate(col_str)) - 1
                t = c.get("t", "")
                v_elem = c.find(f"{{{_NS}}}v")
                if v_elem is None:
                    continue
                raw = v_elem.text or ""
                if t == "s":
                    cells[col_idx] = shared_strings[int(raw)] if raw.isdigit() else raw
                elif t == "b":
                    cells[col_idx] = bool(int(raw))
                else:
                    try:
                        cells[col_idx] = int(raw) if "." not in raw else float(raw)
                    except (ValueError, TypeError):
                        cells[col_idx] = raw
            if cells:
                max_col = max(cells) + 1
                yield [cells.get(j) for j in range(max_col)]


def extract_pricing_spreadsheet(filepath):
    """
    Extract line items from a DHS LLSM pricing XLSX spreadsheet.

    Column layout (0-indexed):
        0=SOW Section, 1=Description, 2=Part# (always empty),
        3=Unit Cost P1, 4=Est Qty P1, 5=Sub-Total P1,
        6=Unit Cost P2, 7=Est Qty P2, 8=Sub-Total P2, ...

    Returns list of dicts with keys:
        sow_section, description, part_number, manufacturer_ref, unit,
        qty_period_1, quantities_by_period, qty, unit_price.
    """
    def _to_qty(v):
        if v is None:
            return None
        try:
            return int(v)
        except (ValueError, TypeError):
            return None

    items = []
    for row_vals in _iter_xlsx_rows(filepath):
        if not row_vals:
            continue
        a = row_vals[0] if len(row_vals) > 0 else None
        if not isinstance(a, str):
            continue
        a_str = a.strip()
        # Must be a terminal item section: exactly 3 numeric parts (e.g. "4.1.1")
        if not re.match(r'^\d+\.\d+\.\d+$', a_str):
            continue

        def cell(idx):
            return row_vals[idx] if idx < len(row_vals) else None

        desc_raw = cell(1)
        desc = re.sub(r'\s{2,}', ' ', str(desc_raw).strip()) if desc_raw else ""

        q1 = _to_qty(cell(4))
        q2 = _to_qty(cell(7))
        q3 = _to_qty(cell(10))
        q4 = _to_qty(cell(13))
        q5 = _to_qty(cell(16))

        items.append({
            "sow_section": a_str,
            "description": desc,
            "part_number": None,
            "manufacturer_ref": None,
            "unit": "EA",
            "qty_period_1": q1 if q1 is not None else "N/A",
            "quantities_by_period": {
                "period_1": q1,
                "period_2": q2,
                "period_3": q3,
                "period_4": q4,
                "period_5": q5,
            },
            "qty": q1 if q1 is not None else "N/A",
            "unit_price": "N/A",
        })

    return items


def merge_line_item_sources(sow_items, pricing_items):
    """
    Merge SOW line items with XLSX pricing items.

    XLSX provides the complete baseline (all items with quantities).
    SOW enriches each item with clean description, manufacturer_ref, part_number, unit.
    """
    sow_by_section = {item["sow_section"]: item for item in sow_items}
    xlsx_sections = set()
    merged = []

    for xlsx_item in pricing_items:
        section = xlsx_item["sow_section"]
        xlsx_sections.add(section)
        merged_item = dict(xlsx_item)

        sow = sow_by_section.get(section)
        if sow:
            if sow["description"] and len(sow["description"]) > 3:
                merged_item["description"] = sow["description"]
            merged_item["manufacturer_ref"] = sow.get("manufacturer_ref")
            merged_item["part_number"] = sow.get("part_number")
            merged_item["unit"] = sow.get("unit", "EA")
            merged_item["spec_text"] = sow.get("spec_text")
            merged_item["source_page"] = sow.get("source_page")
            merged_item["source_file"] = sow.get("source_file")

        merged.append(merged_item)

    # Add any SOW items not in XLSX (guard — should not happen with this fixture)
    for sow_item in sow_items:
        if sow_item["sow_section"] not in xlsx_sections:
            merged.append({
                **sow_item,
                "qty_period_1": "N/A",
                "quantities_by_period": {
                    "period_1": None, "period_2": None,
                    "period_3": None, "period_4": None, "period_5": None,
                },
                "qty": "N/A",
            })

    def _section_key(item):
        return tuple(int(p) for p in item["sow_section"].split(".") if p.isdigit())

    for item in merged:
        has_sow = bool(item.get("spec_text"))
        has_xlsx = item.get("qty_period_1") not in (None, "N/A")
        if has_sow and has_xlsx:
            item["_source"] = "SOW+XLSX"
        elif has_sow:
            item["_source"] = "SOW"
        elif has_xlsx:
            item["_source"] = "XLSX"
        else:
            item["_source"] = "unknown"

        periods = item.get("quantities_by_period", {})
        period_values = [v for v in periods.values() if isinstance(v, (int, float))]
        total = sum(period_values) if period_values else None
        item["qty_total"] = total
        if total is not None:
            item["qty"] = total

    merged.sort(key=_section_key)
    return merged


def extract_line_items(solicitation, text):
    """
    Derive line items from extracted solicitation data and raw text.
    Returns list of dicts: {description, size, unit, qty, unit_price}.
    Any field that cannot be determined is set to the string "N/A".
    """
    base_desc = (solicitation.get("project_title") or
                 solicitation.get("solicitation_number") or "N/A")
    items = []

    # 0. Use AI-extracted line items if available (best quality)
    ai_items = solicitation.get("ai_line_items", [])
    if ai_items:
        for it in ai_items:
            items.append({
                "description": it.get("description", base_desc) or base_desc,
                "size":        it.get("size", "N/A") or "N/A",
                "unit":        it.get("unit", "EA") or "EA",
                "qty":         it.get("qty", "N/A"),
                "unit_price":  it.get("unit_price", "N/A"),
            })
        return items

    # 1. Use size/qty pairs already pulled by the rules engine
    quantities = solicitation.get("quantities", [])
    if quantities:
        for q in quantities:
            try:
                qty_val = int(q.get("qty", "N/A"))
            except (ValueError, TypeError):
                qty_val = "N/A"
            items.append({
                "description": base_desc,
                "size":        q.get("size", "N/A") or "N/A",
                "unit":        "EA",
                "qty":         qty_val,
                "unit_price":  "N/A",
            })
        return items

    # 2. Look for CLIN-style line items
    clin_blocks = re.findall(
        r"(?:CLIN|LINE\s*ITEM|ITEM)\s*(\d{1,4})\s+(.*?)(?=(?:CLIN|LINE\s*ITEM|ITEM)\s*\d|\Z)",
        text, re.IGNORECASE | re.DOTALL
    )
    if clin_blocks:
        for _clin_num, block in clin_blocks:
            block = re.sub(r"\s+", " ", block).strip()
            qty_m = re.search(r"(?:QTY|Quantity)[:\s]*(\d+)", block, re.IGNORECASE)
            up_m  = re.search(r"(?:Unit\s+Price|UNIT\s+PRICE|UP)[:\s]*\$?([\d,\.]+)", block, re.IGNORECASE)
            qty_val = int(qty_m.group(1)) if qty_m else "N/A"
            up_val  = float(up_m.group(1).replace(",","")) if up_m else "N/A"
            items.append({
                "description": block[:120] or base_desc,
                "size":        "N/A",
                "unit":        "EA",
                "qty":         qty_val,
                "unit_price":  up_val,
            })
        return items

    # 3. Fallback: single row from project title
    items.append({"description": base_desc, "size": "N/A", "unit": "EA", "qty": "N/A", "unit_price": "N/A"})
    return items


# ── MULTI-FILE BUNDLE PARSING ─────────────────────────────────────────────────

def classify_document(text, filename):
    """
    Classify a document as 'main', 'sow', or 'pricing' based on filename and content.
    """
    fn = filename.lower()
    if fn.endswith(('.xlsx', '.xls', '.csv')):
        return "pricing"
    if 'pricing' in fn or 'attachment 2' in fn.replace('_', ' ').replace('-', ' '):
        return "pricing"
    if 'sow' in fn or 'statement of work' in fn.replace('_', ' '):
        return "sow"
    if 'Statement of Work' in text[:500]:
        return "sow"
    fmt = detect_format(text)
    if fmt != 'unknown':
        return "main"
    if re.search(r'4\.\d+\.\d+\s+[A-Z]', text[:5000]):
        return "sow"
    return "main"


def _extract_clin_items(text: str) -> list:
    """
    Extract CLIN-numbered line items from SF-1449 and similar DoD solicitations.
    Used as fallback when extract_sow_line_items() returns no 4.x.x items.
    Returns list of dicts compatible with the line_items schema.
    """
    blocks = re.findall(
        r"(?:^|\n)(?:CLIN\s*|ITEM\s*(?:NO\.?\s*)?)(\d{4})\s+"
        r"(.*?)(?=(?:\n(?:CLIN\s*|ITEM\s*(?:NO\.?\s*)?)\d{4}\s)|\Z)",
        text, re.IGNORECASE | re.DOTALL
    )
    if not blocks:
        return []

    items = []
    for clin_num, body in blocks:
        body_clean = re.sub(r"\s+", " ", body).strip()

        desc_m = re.match(r"([^\n.]{5,120})", body_clean)
        description = desc_m.group(1).strip() if desc_m else body_clean[:120]

        unit_m = re.search(r"\b(Each|EA|Pound|LB|Case|Spool|Box)\b", body_clean, re.IGNORECASE)
        unit = unit_m.group(1).upper() if unit_m else "EA"
        if unit == "POUND":
            unit = "LB"

        items.append({
            "sow_section":      f"CLIN {clin_num}",
            "description":      description,
            "manufacturer_ref": None,
            "part_number":      None,
            "unit":             unit,
            "qty":              "N/A",
            "unit_price":       "N/A",
            "spec_text":        body_clean[:2000],
            "source_page":      None,
            "source_file":      None,
            "_source":          "CLIN",
            "qty_total":        None,
        })

    return items


# ── CONFIDENCE SCORING ───────────────────────────────────────────────────────

# Confidence weight constants — adjustable without touching logic
_CONF_WEIGHT_FORMAT = 0.3
_CONF_WEIGHT_FIELDS = 0.4
_CONF_WEIGHT_ITEMS  = 0.3


def compute_confidence(result: dict) -> dict:
    """
    Compute a structured parse-quality confidence score.
    Called at end of parse_solicitation_bundle() before return.
    Do not call from extract_data() — line item counts are not
    available there.
    """

    # --- format_detection score ---
    # Determined from _format and extraction_warnings — no changes
    # to detect_format() return type required.
    known_formats = {"sf1449", "sam_export", "agency_form", "formal_rfq"}
    fmt = result.get("_format", "unknown")
    warnings = result.get("extraction_warnings", [])
    has_unknown_warning = any(w.get("code") == "unknown_format" for w in warnings)
    reasons = []

    if fmt in known_formats:
        format_score = 1.0
        reasons.append(f"Format recognized as {fmt}")
    elif has_unknown_warning:
        format_score = 0.0
        reasons.append("Format not recognized — extraction may be unreliable")
    else:
        format_score = 0.5
        reasons.append("Format uncertain — using fallback extraction")

    # --- required_fields score ---
    required = ["solicitation_number", "due_date", "contact_email", "naics_code"]
    field_labels = {
        "solicitation_number": "solicitation number",
        "due_date":            "due date",
        "contact_email":       "contact email",
        "naics_code":          "NAICS code",
    }
    filled = sum(
        1 for f in required
        if result.get(f) and str(result.get(f)).strip()
    )
    fields_score = filled / len(required)
    missing_labels = [
        field_labels[f] for f in required
        if not (result.get(f) and str(result.get(f)).strip())
    ]
    if not missing_labels:
        reasons.append("All required fields found")
    else:
        reasons.append(f"Missing required fields: {', '.join(missing_labels)}")

    # --- line_items score ---
    items = result.get("line_items", [])
    if not items:
        items_score = 0.0
        reasons.append("No line items found")
    elif any(i.get("_source") == "SOW+XLSX" for i in items):
        items_score = 1.0
        reasons.append(f"Line items: SOW spec and pricing quantities merged ({len(items)} items)")
    elif any(i.get("_source") in ("SOW", "XLSX") for i in items):
        items_score = 0.7
        reasons.append(f"Line items: SOW spec only — no pricing spreadsheet ({len(items)} items)")
    elif any(i.get("_source") == "CLIN" for i in items):
        items_score = 0.4
        reasons.append(f"Line items: CLIN fallback — {len(items)} items extracted from body text")
    elif (
        len(items) == 1 and
        items[0].get("description") in (
            result.get("solicitation_number", ""),
            result.get("project_title", "")
        )
    ):
        items_score = 0.1
        reasons.append("Line items: single placeholder row from project title")
    else:
        items_score = 0.5
        reasons.append(f"Line items: {len(items)} items from unverified source")

    # --- overall (weighted average) ---
    overall = round(
        _CONF_WEIGHT_FORMAT * format_score +
        _CONF_WEIGHT_FIELDS * fields_score +
        _CONF_WEIGHT_ITEMS  * items_score,
        2
    )

    return {
        "overall":          overall,
        "format_detection": format_score,
        "required_fields":  fields_score,
        "line_items":       items_score,
        "warnings":         warnings,
        "reasons":          reasons,
    }


def parse_solicitation_bundle(files):
    """
    Parse a bundle of uploaded files and merge results.

    files: list of dicts with keys: path (str), filename (str)

    1. Load each file with document_loader.load_document()
    2. Classify each as 'main', 'sow', or 'pricing' using classify_document()
    3. Extract header fields from the main doc using extract_data()
    4. Extract SOW line items from sow doc(s) using extract_sow_line_items()
    5. Extract pricing from xlsx doc(s) using extract_pricing_spreadsheet()
    6. Merge with merge_line_item_sources(sow_items, pricing_items)
    7. If no dedicated SOW doc found, try extract_sow_line_items() on main doc text
    8. Set result['line_items'] = merged items
    9. Return result dict
    """
    from document_loader import load_document

    docs = []
    for f in files:
        result = load_document(f["path"])
        role = classify_document(result.text, f["filename"])
        print(f"[parse_solicitation_bundle] {f['filename']} -> role={role}, chars={len(result.text)}")
        docs.append({"role": role, "result": result, "filename": f["filename"], "path": f["path"]})

    # Find main doc — first classified as main, or fall back to first doc overall
    main_doc = next((d for d in docs if d["role"] == "main"), None)
    if main_doc is None:
        main_doc = docs[0] if docs else None

    if main_doc is None:
        return {}

    # Extract header fields from main doc
    data = extract_data(main_doc["result"].text)

    # Extract SOW line items from dedicated sow doc(s)
    sow_docs = [d for d in docs if d["role"] == "sow"]
    sow_items = []
    for sow_doc in sow_docs:
        items = extract_sow_line_items(sow_doc["result"].text, page_texts=sow_doc["result"].page_texts)
        for item in items:
            item["source_file"] = sow_doc["filename"]
        sow_items.extend(items)

    # If no dedicated SOW doc found, try on main doc text
    if not sow_items:
        items = extract_sow_line_items(main_doc["result"].text, page_texts=main_doc["result"].page_texts)
        for item in items:
            item["source_file"] = main_doc["filename"]
        sow_items = items

    # Extract pricing from xlsx doc(s)
    pricing_docs = [d for d in docs if d["role"] == "pricing"]
    pricing_items = []
    for pricing_doc in pricing_docs:
        pricing_items.extend(extract_pricing_spreadsheet(pricing_doc["path"]))

    # Merge and attach to result
    if sow_items or pricing_items:
        data["line_items"] = merge_line_item_sources(sow_items, pricing_items)
        print(f"[parse_solicitation_bundle] line_items={len(data['line_items'])} "
              f"(sow={len(sow_items)}, pricing={len(pricing_items)})")

    # CLIN fallback — fires only when no 4.x.x SOW items and no pricing XLSX
    if not data.get("line_items") and main_doc:
        clin_items = _extract_clin_items(main_doc["result"].text)
        if clin_items:
            data["line_items"] = clin_items
            print(f"[parse_solicitation_bundle] CLIN fallback: {len(clin_items)} items")

    # ── assemble extraction_warnings ──────────────────────────────────────────────
    warnings = []
    for _field in ["solicitation_number", "due_date", "contact_email", "naics_code"]:
        _val = data.get(_field)
        if not _val or str(_val).strip() == "":
            warnings.append({"code": "missing_field", "field": _field})
    if data.get("_format") == "unknown":
        warnings.append({"code": "unknown_format"})
    if not data.get("line_items"):
        warnings.append({"code": "no_line_items", "source": "fallback_single_row"})
    data["extraction_warnings"] = warnings
    if warnings:
        print(f"[parse_solicitation_bundle] warnings={warnings}")

    data["confidence"] = compute_confidence(data)
    return data
