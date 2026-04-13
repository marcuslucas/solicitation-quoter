#!/usr/bin/env python3
"""
document_loader.py — Magic-byte-aware document loader for SolicitationQuoter.

Dispatches to the right extraction strategy based on actual file content,
not file extension. SAM.gov delivers some solicitations as ZIP archives with
a .pdf extension; this module handles that transparently.
"""

import json
import zipfile
from dataclasses import dataclass, field
from pathlib import Path


@dataclass
class DocumentResult:
    text: str
    """Reassembled full text of the document, ready for regex extraction."""

    page_count: int = 0
    """Number of pages (or ZIP text entries) extracted."""

    source_format: str = "unknown"
    """Detected container format: 'pdf', 'sam_zip', 'docx', 'txt', 'unknown'."""

    has_images: bool = False
    """True if the source contains image files (JPEG/PNG pages in a ZIP)."""

    page_texts: list = field(default_factory=list)
    """Per-page text strings. Index 0 = page 1. Enables targeted extraction."""

    error: str = None
    """Non-None if loading failed; text may be empty or partial."""


def load_document(filepath: str) -> DocumentResult:
    """
    Detect the actual file type by magic bytes and extract text.

    Dispatch table:
      b'PK\\x03\\x04'  -> _load_zip_bundle()  (SAM.gov ZIP or DOCX)
      b'%PDF'          -> _load_pdf()
      .txt extension   -> _load_text()
      anything else    -> try _load_pdf(), then return error result

    Never raises; all errors are returned in DocumentResult.error.
    Magic bytes take absolute precedence over file extension — a file
    named .pdf that starts with PK is a ZIP, not a PDF.
    """
    path = Path(filepath)

    try:
        with open(filepath, "rb") as f:
            magic = f.read(4)
    except OSError as e:
        return DocumentResult(text="", error=f"Cannot open file: {e}")

    # ZIP-based formats (SAM.gov bundles and DOCX both start with PK)
    if magic[:2] == b"PK":
        return _load_zip_bundle(filepath)

    # Real PDF
    if magic[:4] == b"%PDF":
        return _load_pdf(filepath)

    # Plain text by extension
    if path.suffix.lower() == ".txt":
        return _load_text(filepath)

    # Unknown — try PDF as last resort
    result = _load_pdf(filepath)
    if result.text:
        return result
    return DocumentResult(
        text="",
        error=f"Unrecognized file format (magic={magic!r}, suffix={path.suffix!r})",
    )


def _load_pdf(filepath: str) -> DocumentResult:
    """
    Extract text from a real PDF using pdfplumber (primary) then pypdf (fallback).

    Preserves per-page text in page_texts for downstream targeted extraction.
    Returns source_format='pdf'.
    On total failure returns DocumentResult(text='', error=<message>).
    """
    page_texts = []

    # Primary: pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_texts.append(page.extract_text() or "")
        text = "\n\n".join(page_texts).strip()
        if text:
            return DocumentResult(
                text=text,
                page_count=len(page_texts),
                source_format="pdf",
                page_texts=page_texts,
            )
    except Exception as e:
        print(f"pdfplumber failed: {e}")

    # Fallback: pypdf
    page_texts = []
    try:
        from pypdf import PdfReader
        reader = PdfReader(filepath)
        for page in reader.pages:
            page_texts.append(page.extract_text() or "")
        text = "\n\n".join(page_texts).strip()
        if text:
            return DocumentResult(
                text=text,
                page_count=len(page_texts),
                source_format="pdf",
                page_texts=page_texts,
            )
    except Exception as e:
        print(f"pypdf failed: {e}")

    return DocumentResult(text="", error="PDF extraction failed with both pdfplumber and pypdf")


def _load_zip_bundle(filepath: str) -> DocumentResult:
    """
    Handle ZIP archives that may be:
      (a) SAM.gov bundles: contain manifest.json + N.txt + N.jpeg per page
      (b) DOCX files: contain word/document.xml

    For SAM.gov bundles:
      - Reads manifest.json to get num_pages
      - Reads 1.txt, 2.txt, ... num_pages.txt in order
      - Joins with double-newline separator to preserve page breaks
      - Sets source_format='sam_zip', has_images=True if any *.jpeg present

    For DOCX:
      - Delegates to _load_docx(filepath)

    For unrecognized ZIPs:
      - Returns DocumentResult(text='', error='ZIP archive without recognized structure')

    If manifest.json is present but a page text file is missing, skips that
    page and logs the gap with print().
    """
    try:
        with zipfile.ZipFile(filepath) as zf:
            names = zf.namelist()

            # SAM.gov bundle: has manifest.json with page metadata
            if "manifest.json" in names:
                manifest = json.loads(zf.read("manifest.json").decode("utf-8", errors="ignore"))
                num_pages = manifest.get("num_pages", 0)
                page_texts = []

                for i in range(1, num_pages + 1):
                    txt_name = f"{i}.txt"
                    if txt_name in names:
                        page_text = zf.read(txt_name).decode("utf-8", errors="ignore")
                        page_texts.append(page_text)
                    else:
                        print(f"[document_loader] Missing page text: {txt_name} (skipping)")

                has_images = any(
                    n.lower().endswith((".jpeg", ".jpg", ".png")) for n in names
                )
                full_text = "\n\n".join(page_texts).strip()

                return DocumentResult(
                    text=full_text,
                    page_count=num_pages,
                    source_format="sam_zip",
                    has_images=has_images,
                    page_texts=page_texts,
                )

            # DOCX file: has word/document.xml
            if "word/document.xml" in names:
                return _load_docx(filepath)

    except zipfile.BadZipFile as e:
        return DocumentResult(text="", error=f"Invalid ZIP archive: {e}")
    except Exception as e:
        return DocumentResult(text="", error=f"ZIP extraction failed: {e}")

    return DocumentResult(text="", error="ZIP archive without recognized structure")


def _load_docx(filepath: str) -> DocumentResult:
    """
    Extract text from a DOCX file using python-docx.
    Extracts paragraph text and table cell text (pipe-separated rows).
    Returns source_format='docx'.
    """
    try:
        from docx import Document
        doc = Document(filepath)
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts)
        return DocumentResult(
            text=text,
            page_count=1,
            source_format="docx",
            page_texts=[text],
        )
    except Exception as e:
        return DocumentResult(text="", error=f"DOCX extraction failed: {e}")


def _load_text(filepath: str) -> DocumentResult:
    """
    Read a plain-text file as UTF-8 (errors='ignore').
    Returns source_format='txt', page_count=1.
    """
    try:
        text = open(filepath, encoding="utf-8", errors="ignore").read()
        return DocumentResult(
            text=text,
            page_count=1,
            source_format="txt",
            page_texts=[text],
        )
    except Exception as e:
        return DocumentResult(text="", error=f"Text file read failed: {e}")
