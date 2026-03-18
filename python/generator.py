#!/usr/bin/env python3
"""DOCX quote generation for SolicitationQuoter."""

import io
import datetime


def generate_quote(solicitation, vendor, line_items):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(1); sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1.25); sec.right_margin = Inches(1.25)

    NAVY  = RGBColor(0x00,0x00,0x00)
    DGRAY = RGBColor(0x33,0x33,0x33)
    WHITE = RGBColor(0xFF,0xFF,0xFF)

    def bg(cell, color):
        tc = cell._tc; pr = tc.get_or_add_tcPr()
        s = OxmlElement("w:shd")
        s.set(qn("w:fill"), color); s.set(qn("w:val"), "clear")
        pr.append(s)

    def run(para, text, bold=False, size=11, color=None, italic=False):
        r = para.add_run(text); r.bold=bold; r.italic=italic; r.font.size=Pt(size)
        r.font.name = 'Calibri'
        if color: r.font.color.rgb=color
        return r

    def heading(txt, sb=14):
        p = doc.add_paragraph()
        p.paragraph_format.space_before=Pt(sb); p.paragraph_format.space_after=Pt(4)
        p.paragraph_format.keep_with_next = True
        run(p, txt, bold=True, size=11, color=NAVY)
        pr = p._p.get_or_add_pPr(); bd = OxmlElement("w:pBdr")
        b = OxmlElement("w:bottom"); b.set(qn("w:val"),"single"); b.set(qn("w:sz"),"6")
        b.set(qn("w:space"),"1"); b.set(qn("w:color"),"1A1A1A"); bd.append(b); pr.append(bd)

    def row_keep(row, header=False):
        tr = row._tr
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = OxmlElement('w:trPr'); tr.insert(0, trPr)
        cant = OxmlElement('w:cantSplit'); cant.set(qn('w:val'), '1')
        trPr.append(cant)
        if header:
            th = OxmlElement('w:tblHeader'); th.set(qn('w:val'), '1')
            trPr.append(th)

    def fmt_num(v):
        """Return float for a numeric value, or None if it is N/A / blank / unparseable."""
        if v is None or v == "" or v == "N/A": return None
        try:
            f = float(v)
            return None if (f != f) else f  # reject NaN
        except (ValueError, TypeError): return None

    def no_borders(table):
        tbl = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
        for existing in tblPr.findall(qn('w:tblBorders')):
            tblPr.remove(existing)
        tblBorders = OxmlElement('w:tblBorders')
        for side in ('top','left','bottom','right','insideH','insideV'):
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'none'); el.set(qn('w:sz'), '0'); el.set(qn('w:color'), 'auto')
            tblBorders.append(el)
        tblPr.append(tblBorders)

    today = datetime.date.today().strftime("%B %d, %Y")

# Header table
    ht = doc.add_table(rows=1, cols=2); ht.style="Table Grid"; ht.autofit=False
    ht.columns[0].width=Inches(3.0); ht.columns[1].width=Inches(3.0)
    ht.alignment = WD_TABLE_ALIGNMENT.CENTER
    lc=ht.cell(0,0); rc=ht.cell(0,1)
    bg(lc,"EFEFEF"); bg(rc,"F5F5F5")
    no_borders(ht)

    # Left cell — company info (logo if provided)
    logo_b64 = vendor.get("logo_b64","")
    if logo_b64:
        try:
            import base64 as _b64
            logo_bytes = _b64.b64decode(logo_b64)
            logo_stream = io.BytesIO(logo_bytes)
            logo_p = lc.paragraphs[0]
            logo_p.paragraph_format.space_before = Pt(8)
            logo_p.paragraph_format.left_indent = Pt(10)
            logo_p.add_run().add_picture(logo_stream, width=Inches(1.4))
            lp = lc.add_paragraph()
        except Exception as _e:
            print(f"Logo insert failed: {_e}")
            lp = lc.paragraphs[0]
    else:
        lp = lc.paragraphs[0]
    lp.paragraph_format.space_before=Pt(10)
    lp.paragraph_format.left_indent=Pt(10)
    run(lp, vendor.get("company_name","Your Company"), bold=True, size=15, color=NAVY)
    for f in ["address","city_state_zip","phone","email","website"]:
        val=vendor.get(f,"")
        if val:
            fp=lc.add_paragraph()
            fp.paragraph_format.left_indent=Pt(10)
            run(fp, val, size=9, color=DGRAY)
    pad=lc.add_paragraph()
    pad.paragraph_format.left_indent=Pt(10)
    run(pad, " ", size=9, color=DGRAY)

    # Right cell — quote metadata (dark text on light background)
    rp=rc.paragraphs[0]
    rp.paragraph_format.space_before=Pt(10)
    rp.paragraph_format.left_indent=Pt(8)
    run(rp,"QUOTE / PROPOSAL", bold=True, size=13, color=RGBColor(0x00,0x00,0x00))
    for label,value in [
        ("Quote #", vendor.get("quote_number","Q-"+datetime.date.today().strftime("%Y%m%d"))),
        ("Date", today),
        ("Solicitation #", solicitation.get("solicitation_number","")),
        ("Valid For", vendor.get("validity_period","30 days")),
        ("Delivery", f"{vendor.get('delivery_days','')} days ARO" if vendor.get("delivery_days") else ""),
        ("Prepared By", vendor.get("prepared_by","")),
    ]:
        if value:
            mp=rc.add_paragraph()
            mp.paragraph_format.left_indent=Pt(8)
            run(mp,f"{label}: ",bold=True,size=9,color=RGBColor(0x33,0x33,0x33))
            run(mp,value,size=9,color=RGBColor(0x33,0x33,0x33))
    rc.add_paragraph()
    doc.add_paragraph()
    # Solicitation info
    heading("SOLICITATION INFORMATION")
    fields=[
        ("Project / Subject",    solicitation.get("project_title")),
        ("Solicitation Number",  solicitation.get("solicitation_number")),
        ("Solicitation Type",    solicitation.get("solicitation_type")),
        ("Issuing Agency",       solicitation.get("issuing_agency")),
        ("Contracting Office",   solicitation.get("contracting_office_address")),
        ("Response Due Date",    solicitation.get("due_date")),
        ("NAICS Code",           solicitation.get("naics_code")),
        ("PSC Code",             solicitation.get("psc_code")),
        ("Set-Aside",            solicitation.get("set_aside")),
        ("Place of Performance", solicitation.get("place_of_performance")),
        ("Period of Performance",solicitation.get("period_of_performance")),
    ]
    fields=[(l,v) for l,v in fields if v]
    if fields:
        st=doc.add_table(rows=len(fields),cols=2); st.style="Table Grid"; st.autofit=False
        st.columns[0].width=Inches(2.0); st.columns[1].width=Inches(4.0)
        st.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i,(label,value) in enumerate(fields):
            row_keep(st.rows[i])
            lc2=st.cell(i,0); rc2=st.cell(i,1)
            bg(lc2,"F0F0F0")
            lp2=lc2.paragraphs[0]; lp2.paragraph_format.left_indent=Pt(4)
            run(lp2,label,bold=True,size=9,color=NAVY)
            rp2=rc2.paragraphs[0]; rp2.paragraph_format.left_indent=Pt(4)
            run(rp2,str(value)[:200],size=9,color=DGRAY)

    # Scope
    scope=solicitation.get("scope_of_work","")
    if scope:
        heading("SCOPE OF WORK / REQUIREMENT")
        sp=doc.add_paragraph(); sp.paragraph_format.left_indent=Inches(0.1)
        run(sp, scope[:2000], size=9.5, color=DGRAY)

    # Line items
    heading("QUOTE DETAILS")
    cw=[Inches(0.35),Inches(2.1),Inches(0.55),Inches(0.6),Inches(0.55),Inches(0.85),Inches(1.0)]
    hdrs=["#","Description / Item","Size/Type","UOM","Qty","Unit Price","Total"]
    lt=doc.add_table(rows=1+len(line_items)+1,cols=7); lt.style="Table Grid"; lt.autofit=False
    lt.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci,w in enumerate(cw): lt.columns[ci].width=w
    hr=lt.rows[0]
    row_keep(hr, header=True)
    for ci,h in enumerate(hdrs):
        c=hr.cells[ci]; bg(c,"1A1A1A"); c.width=cw[ci]
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run(p,h,bold=True,size=9,color=WHITE)
    grand=0.0; has_any_price=False
    AL=[WD_ALIGN_PARAGRAPH.CENTER,WD_ALIGN_PARAGRAPH.LEFT,WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,WD_ALIGN_PARAGRAPH.CENTER,WD_ALIGN_PARAGRAPH.RIGHT,
        WD_ALIGN_PARAGRAPH.RIGHT]
    for i,item in enumerate(line_items):
        row=lt.rows[i+1]; row_keep(row)
        qty_n = fmt_num(item.get("qty"))
        up_n  = fmt_num(item.get("unit_price"))
        total_n = (qty_n * up_n) if (qty_n is not None and up_n is not None) else None
        if total_n is not None: grand += total_n; has_any_price = True
        bcolor="FFFFFF" if i%2==0 else "F7F9FC"
        qty_s   = (str(int(qty_n)) if qty_n is not None and qty_n == int(qty_n) else str(qty_n)) if qty_n is not None else "N/A"
        up_s    = f"${up_n:,.2f}" if up_n is not None else "N/A"
        total_s = f"${total_n:,.2f}" if total_n is not None else "N/A"
        desc    = item.get("description","") or "N/A"
        size    = item.get("size","") or "N/A"
        unit    = item.get("unit","EA") or "EA"
        vals=[str(i+1), desc, size, unit, qty_s, up_s, total_s]
        for ci,(val,al,w) in enumerate(zip(vals,AL,cw)):
            c=row.cells[ci]; bg(c,bcolor); c.width=w
            p=c.paragraphs[0]; p.alignment=al; p.paragraph_format.left_indent=Pt(3)
            run(p,val,size=9,color=DGRAY)
    tr=lt.rows[-1]; row_keep(tr)
    for ci in range(7): bg(tr.cells[ci],"F0F0F0"); tr.cells[ci].width=cw[ci]
    tr.cells[0].merge(tr.cells[4])
    tp=tr.cells[0].paragraphs[0]; tp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    run(tp,"GRAND TOTAL",bold=True,size=10,color=NAVY)
    freight=float(vendor.get("freight",0) or 0)
    tax_rate=float(vendor.get("tax_rate",0) or 0)
    tax=grand*(tax_rate/100); final=grand+freight+tax
    gp=tr.cells[-1].paragraphs[0]; gp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    final_s = f"${final:,.2f}" if (has_any_price or freight>0) else "N/A"
    run(gp,final_s,bold=True,size=11,color=NAVY)

    # Notes & Terms
    notes=vendor.get("notes",""); terms=vendor.get("terms","")
    if notes or terms:
        heading("NOTES & TERMS")
        if notes:
            np=doc.add_paragraph(); np.paragraph_format.left_indent=Inches(0.1)
            run(np,"Notes: ",bold=True,size=9.5,color=NAVY); run(np,notes,size=9.5,color=DGRAY)
        if terms:
            tp2=doc.add_paragraph(); tp2.paragraph_format.left_indent=Inches(0.1)
            run(tp2,"Terms: ",bold=True,size=9.5,color=NAVY); run(tp2,terms,size=9.5,color=DGRAY)

    # Option years summary
    option_years = vendor.get("option_years", [])
    if vendor.get("option_years_enabled") and option_years and has_any_price:
        base_total = grand + freight + tax
        heading("OPTION YEAR PRICING SUMMARY")
        oy_rows = [("Base Year (Year 1)", base_total)]
        for oy in option_years:
            pct = float(oy.get("pct", 0) or 0)
            oy_rows.append((oy.get("label","Option Year"), base_total * (1 + pct / 100)))
        oy_rows.append(("TOTAL — All Periods", sum(r[1] for r in oy_rows)))
        ot = doc.add_table(rows=len(oy_rows), cols=2); ot.style="Table Grid"; ot.autofit=False
        ot.columns[0].width = Inches(3.5); ot.columns[1].width = Inches(2.5)
        ot.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i,(period,total_oy) in enumerate(oy_rows):
            row_keep(ot.rows[i])
            is_last = (i == len(oy_rows) - 1)
            lc3=ot.cell(i,0); rc3=ot.cell(i,1)
            bg(lc3,"F0F0F0" if (i==0 or is_last) else "FFFFFF")
            bg(rc3,"F0F0F0" if (i==0 or is_last) else "FFFFFF")
            lp3=lc3.paragraphs[0]; lp3.paragraph_format.left_indent=Pt(4)
            run(lp3,period,bold=is_last,size=9,color=NAVY)
            rp3=rc3.paragraphs[0]; rp3.alignment=WD_ALIGN_PARAGRAPH.RIGHT
            run(rp3,f"${total_oy:,.2f}",bold=is_last,size=9,color=NAVY if is_last else DGRAY)

    # Signature
    heading("AUTHORIZED SIGNATURE", sb=18)
    sigt=doc.add_table(rows=1,cols=2); sigt.style="Table Grid"; sigt.autofit=False
    sigt.columns[0].width=Inches(3.0); sigt.columns[1].width=Inches(3.0)
    sigt.alignment = WD_TABLE_ALIGNMENT.CENTER
    row_keep(sigt.rows[0])
    lsc=sigt.cell(0,0); rsc=sigt.cell(0,1)
    def sigline(cell,label,value=""):
        p=cell.add_paragraph(); p.paragraph_format.left_indent=Pt(6); p.paragraph_format.space_after=Pt(10)
        run(p,f"{label}: ",bold=True,size=9,color=NAVY); run(p,value or "_"*28,size=9,color=DGRAY)
    sigline(lsc,"Authorized Signature"); sigline(lsc,"Printed Name",vendor.get("prepared_by",""))
    sigline(lsc,"Title",vendor.get("title","")); sigline(lsc,"Date",today)
    sigline(rsc,"Company",vendor.get("company_name","")); sigline(rsc,"Phone",vendor.get("phone",""))
    sigline(rsc,"Email",vendor.get("email","")); sigline(rsc,"SAM UEI",vendor.get("sam_uei",""))
    doc.add_paragraph()
    fp=doc.add_paragraph(); fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run(fp,f"Quote valid for {vendor.get('validity_period','30 days')} from {today}.",
        size=8,color=RGBColor(0x99,0x99,0x99),italic=True)

    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf.read()
