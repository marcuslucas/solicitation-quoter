#!/usr/bin/env python3
import os, sys, json, re, tempfile, traceback, io
from pathlib import Path
from flask import Flask, request, jsonify, send_file

app = Flask(__name__)

@app.after_request
def cors(r):
    r.headers["Access-Control-Allow-Origin"] = "http://127.0.0.1:5199"
    r.headers["Access-Control-Allow-Headers"] = "Content-Type"
    r.headers["Access-Control-Allow-Methods"] = "GET,POST,OPTIONS"
    return r

@app.route("/ping")
def ping():
    return jsonify({"status": "ok"})

# ── PARSERS ──────────────────────────────────────────────────────────────────

def parse_pdf(filepath):
    try:
        import pdfplumber
        with pdfplumber.open(filepath) as pdf:
            text = "\n\n".join(p.extract_text() or "" for p in pdf.pages).strip()
        if text:
            return text
    except Exception as e:
        print(f"pdfplumber failed: {e}")
    try:
        from pypdf import PdfReader
        reader = PdfReader(filepath)
        text = "\n\n".join(p.extract_text() or "" for p in reader.pages).strip()
        if text:
            return text
    except Exception as e:
        print(f"pypdf failed: {e}")
    return ""

def parse_docx(filepath):
    from docx import Document
    doc = Document(filepath)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)

def parse_document(filepath):
    ext = Path(filepath).suffix.lower()
    if ext == ".pdf":    return parse_pdf(filepath)
    if ext in (".docx",".doc"): return parse_docx(filepath)
    if ext == ".txt":
        return open(filepath, encoding="utf-8", errors="ignore").read()
    raise ValueError(f"Unsupported file type: {ext}")

# ── RULE-BASED EXTRACTOR ─────────────────────────────────────────────────────

def extract(text):
    def find(patterns):
        for p in patterns:
            m = re.search(p, text, re.IGNORECASE | re.MULTILINE)
            if m:
                v = (m.group(1) if m.lastindex else m.group(0)).strip()
                if v: return v
        return ""

    d = {}
    d["solicitation_number"] = find([
        r"Notice\s*ID[:\s]*([A-Z0-9\-]+)",
        r"Solicitation\s*(?:No|Number|#)[:\s]*([A-Z0-9\-]+)",
        r"RFP\s*(?:No|Number|#)[:\s]*([A-Z0-9\-]+)",
        r"RFQ\s*(?:No|Number|#)[:\s]*([A-Z0-9\-]+)",
    ])
    d["project_title"] = find([
        r"Subject[:\s]+(.+?)(?:\n|$)",
        r"(?:Project|Contract|Solicitation)\s+Title[:\s]+(.+?)(?:\n|$)",
    ])
    d["solicitation_type"] = find([
        r"Solicitation\s+Type[:\s]+(\w+)",
        r"\b(RFQ|RFP|IFB|Sources Sought|Combined Synopsis)\b",
    ])
    d["issuing_agency"] = find([
        r"Contracting\s+Office\s+Name[:\s]+(.+?)(?:\n|$)",
        r"(?:Issued by|Issuing Office|Agency)[:\s]+(.+?)(?:\n|$)",
    ])
    d["contracting_office_address"] = find([
        r"Contracting\s+Office\s+Address\s*\n(.+?)(?:\n\n|\Z)",
    ])
    d["due_date"] = find([
        r"Response\s+Date[:\s]*([^\n]+)",
        r"(?:Due Date|Deadline|Response Due)[:\s]*([^\n]+)",
    ])
    d["posting_date"] = find([r"Posting\s+Date[:\s]*([^\n]+)"])
    d["contact_name"]  = find([r"Primary\s+Contact\s+Name[:\s]*([^\n]+)"])
    d["contact_email"] = find([
        r"Primary\s+Contact\s+Email[:\s]*([^\s]+@[^\s]+)",
        r"([a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,})",
    ])
    d["contact_phone"] = find([r"Primary\s+Contact\s+Phone\s*(?:Number)?[:\s]*([^\n]+)"])
    d["naics_code"] = find([r"NAICS\s*(\d{5,6}(?:[:\s]+[^\n]+)?)"])
    d["psc_code"]   = find([r"Product\s+or\s+Service\s+Code\s*([0-9A-Z]+[^\n]*)"])
    d["set_aside"]  = find([
        r"Set\s*Aside\s*Code\s*(.+?)(?:\n|$)",
        r"(Total Small Business Set.?Aside|Small Business|8\(a\)|SDVOSB|HUBZone|WOSB)",
    ])
    d["place_of_performance"] = find([
        r"Place\s+of\s+Performance\s+Address\s*([^\n]+(?:\n[^\n]+)?)",
    ])
    d["period_of_performance"] = find([r"Period\s+of\s+Performance[:\s]+(.+?)(?:\n|$)"])
    d["estimated_value"] = find([
        r"Estimated\s+(?:Value|Cost|Budget)[:\s]*\$?([\d,\.]+(?:\s*(?:million|billion))?)",
        r"Not\s+to\s+Exceed[:\s]*\$?([\d,\.]+)",
    ])

    # Scope of work
    m = re.search(r"Description[:\s]*(.+?)(?=Contact Information|Notice Details|Attachment|\Z)", text, re.IGNORECASE|re.DOTALL)
    d["scope_of_work"] = re.sub(r"\s+"," ", m.group(1)).strip()[:3000] if m else re.sub(r"\s+"," ",text[:2000]).strip()

    # Quantities
    qtys = re.findall(r"\b(SM|S|M|L|XL|XXL|2XL|3XL)[:\s]*(\d+)", text, re.IGNORECASE)
    if qtys:
        d["quantities"] = [{"size": q[0].upper(), "qty": q[1]} for q in qtys]

    # Attachments
    atts = re.findall(r"[•\-\*]?\s*([A-Z0-9_\-]+\.pdf|[A-Z0-9_\-]+\.docx)", text, re.IGNORECASE)
    if atts:
        d["attachments"] = list(set(atts))

    return {k: v for k, v in d.items() if v not in ("", [], None)}

# ── AI EXTRACTOR ─────────────────────────────────────────────────────────────

def ai_extract(text, api_key):
    import anthropic
    client = anthropic.Anthropic(api_key=api_key)
    if len(text) > 14000:
        text = text[:7000] + "\n\n[...truncated...]\n\n" + text[-7000:]
    prompt = f"""Extract all data from this solicitation and return ONLY a JSON object with these keys (empty string if not found):
solicitation_number, project_title, solicitation_type, issuing_agency, contracting_office_address,
due_date, posting_date, contact_name, contact_email, contact_phone, naics_code, psc_code, set_aside,
place_of_performance, contract_type, period_of_performance, estimated_value, scope_of_work,
quantities (array of {{size,qty}}), special_requirements (array of strings), attachments (array of strings),
line_items (array of {{description, size, unit, qty, unit_price}} — extract from CLIN tables or item lists if present; use empty string for unknown unit_price)

SOLICITATION:
{text}"""
    resp = client.messages.create(model="claude-sonnet-4-6", max_tokens=2000,
                                   messages=[{"role":"user","content":prompt}])
    raw = re.sub(r"^```(?:json)?\s*","",resp.content[0].text.strip())
    raw = re.sub(r"\s*```$","",raw)
    return json.loads(raw)

def extract_data(text, api_key=""):
    rules = extract(text)
    if api_key:
        try:
            ai = ai_extract(text, api_key)
            merged = {**rules}
            for k,v in ai.items():
                if v and v != "" and v != [] and v != {}:
                    merged[k] = v
            # Store AI line items separately so extractor can prefer them
            if ai.get("line_items"):
                merged["ai_line_items"] = ai["line_items"]
            merged["_method"] = "ai+rules"
            return merged
        except Exception as e:
            print(f"AI failed, using rules: {e}")
    rules["_method"] = "rules"
    return rules

# ── LINE ITEM EXTRACTOR ───────────────────────────────────────────────────────

def extract_line_items(solicitation, text):
    """
    Derive line items from extracted solicitation data and raw text.
    Returns list of dicts: {description, size, unit, qty, unit_price}.
    Any field that cannot be determined is set to the string "N/A".
    """
    base_desc = (solicitation.get("project_title") or
                 solicitation.get("solicitation_number") or "N/A")
    items = []

    # 0. Use AI-extracted line items if available (best quality)
    ai_items = solicitation.get("ai_line_items", [])
    if ai_items:
        for it in ai_items:
            items.append({
                "description": it.get("description", base_desc) or base_desc,
                "size":        it.get("size", "N/A") or "N/A",
                "unit":        it.get("unit", "EA") or "EA",
                "qty":         it.get("qty", "N/A"),
                "unit_price":  it.get("unit_price", "N/A"),
            })
        return items

    # 1. Use size/qty pairs already pulled by the rules engine
    quantities = solicitation.get("quantities", [])
    if quantities:
        for q in quantities:
            try:
                qty_val = int(q.get("qty", "N/A"))
            except (ValueError, TypeError):
                qty_val = "N/A"
            items.append({
                "description": base_desc,
                "size":        q.get("size", "N/A") or "N/A",
                "unit":        "EA",
                "qty":         qty_val,
                "unit_price":  "N/A",
            })
        return items

    # 2. Look for CLIN-style line items
    clin_blocks = re.findall(
        r"(?:CLIN|LINE\s*ITEM|ITEM)\s*(\d{1,4})\s+(.*?)(?=(?:CLIN|LINE\s*ITEM|ITEM)\s*\d|\Z)",
        text, re.IGNORECASE | re.DOTALL
    )
    if clin_blocks:
        for _clin_num, block in clin_blocks:
            block = re.sub(r"\s+", " ", block).strip()
            qty_m = re.search(r"(?:QTY|Quantity)[:\s]*(\d+)", block, re.IGNORECASE)
            up_m  = re.search(r"(?:Unit\s+Price|UNIT\s+PRICE|UP)[:\s]*\$?([\d,\.]+)", block, re.IGNORECASE)
            qty_val = int(qty_m.group(1)) if qty_m else "N/A"
            up_val  = float(up_m.group(1).replace(",","")) if up_m else "N/A"
            items.append({
                "description": block[:120] or base_desc,
                "size":        "N/A",
                "unit":        "EA",
                "qty":         qty_val,
                "unit_price":  up_val,
            })
        return items

    # 3. Fallback: single row from project title
    items.append({"description": base_desc, "size": "N/A", "unit": "EA", "qty": "N/A", "unit_price": "N/A"})
    return items

# ── QUOTE GENERATOR ──────────────────────────────────────────────────────────

def generate_quote(solicitation, vendor, line_items):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import datetime

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(1); sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1.25); sec.right_margin = Inches(1.25)

    NAVY  = RGBColor(0x00,0x00,0x00)
    DGRAY = RGBColor(0x33,0x33,0x33)
    WHITE = RGBColor(0xFF,0xFF,0xFF)

    def bg(cell, color):
        tc = cell._tc; pr = tc.get_or_add_tcPr()
        s = OxmlElement("w:shd")
        s.set(qn("w:fill"), color); s.set(qn("w:val"), "clear")
        pr.append(s)

    def run(para, text, bold=False, size=11, color=None, italic=False):
        r = para.add_run(text); r.bold=bold; r.italic=italic; r.font.size=Pt(size)
        r.font.name = 'Calibri'
        if color: r.font.color.rgb=color
        return r

    def heading(txt, sb=14):
        p = doc.add_paragraph()
        p.paragraph_format.space_before=Pt(sb); p.paragraph_format.space_after=Pt(4)
        p.paragraph_format.keep_with_next = True
        run(p, txt, bold=True, size=11, color=NAVY)
        pr = p._p.get_or_add_pPr(); bd = OxmlElement("w:pBdr")
        b = OxmlElement("w:bottom"); b.set(qn("w:val"),"single"); b.set(qn("w:sz"),"6")
        b.set(qn("w:space"),"1"); b.set(qn("w:color"),"1A1A1A"); bd.append(b); pr.append(bd)

    def row_keep(row, header=False):
        tr = row._tr
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = OxmlElement('w:trPr'); tr.insert(0, trPr)
        cant = OxmlElement('w:cantSplit'); cant.set(qn('w:val'), '1')
        trPr.append(cant)
        if header:
            th = OxmlElement('w:tblHeader'); th.set(qn('w:val'), '1')
            trPr.append(th)

    def fmt_num(v):
        """Return float for a numeric value, or None if it is N/A / blank / unparseable."""
        if v is None or v == "" or v == "N/A": return None
        try:
            f = float(v)
            return None if (f != f) else f  # reject NaN
        except (ValueError, TypeError): return None

    def no_borders(table):
        tbl = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
        for existing in tblPr.findall(qn('w:tblBorders')):
            tblPr.remove(existing)
        tblBorders = OxmlElement('w:tblBorders')
        for side in ('top','left','bottom','right','insideH','insideV'):
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'none'); el.set(qn('w:sz'), '0'); el.set(qn('w:color'), 'auto')
            tblBorders.append(el)
        tblPr.append(tblBorders)

    today = datetime.date.today().strftime("%B %d, %Y")

# Header table
    ht = doc.add_table(rows=1, cols=2); ht.style="Table Grid"; ht.autofit=False
    ht.columns[0].width=Inches(3.0); ht.columns[1].width=Inches(3.0)
    ht.alignment = WD_TABLE_ALIGNMENT.CENTER
    lc=ht.cell(0,0); rc=ht.cell(0,1)
    bg(lc,"EFEFEF"); bg(rc,"F5F5F5")
    no_borders(ht)

    # Left cell — company info (logo if provided)
    logo_b64 = vendor.get("logo_b64","")
    if logo_b64:
        try:
            import base64 as _b64
            logo_bytes = _b64.b64decode(logo_b64)
            logo_stream = io.BytesIO(logo_bytes)
            logo_p = lc.paragraphs[0]
            logo_p.paragraph_format.space_before = Pt(8)
            logo_p.paragraph_format.left_indent = Pt(10)
            logo_p.add_run().add_picture(logo_stream, width=Inches(1.4))
            lp = lc.add_paragraph()
        except Exception as _e:
            print(f"Logo insert failed: {_e}")
            lp = lc.paragraphs[0]
    else:
        lp = lc.paragraphs[0]
    lp.paragraph_format.space_before=Pt(10)
    lp.paragraph_format.left_indent=Pt(10)
    run(lp, vendor.get("company_name","Your Company"), bold=True, size=15, color=NAVY)
    for f in ["address","city_state_zip","phone","email","website"]:
        val=vendor.get(f,"")
        if val:
            fp=lc.add_paragraph()
            fp.paragraph_format.left_indent=Pt(10)
            run(fp, val, size=9, color=DGRAY)
    pad=lc.add_paragraph()
    pad.paragraph_format.left_indent=Pt(10)
    run(pad, " ", size=9, color=DGRAY)

    # Right cell — quote metadata (dark text on light background)
    rp=rc.paragraphs[0]
    rp.paragraph_format.space_before=Pt(10)
    rp.paragraph_format.left_indent=Pt(8)
    run(rp,"QUOTE / PROPOSAL", bold=True, size=13, color=RGBColor(0x00,0x00,0x00))
    for label,value in [
        ("Quote #", vendor.get("quote_number","Q-"+datetime.date.today().strftime("%Y%m%d"))),
        ("Date", today),
        ("Solicitation #", solicitation.get("solicitation_number","")),
        ("Valid For", vendor.get("validity_period","30 days")),
        ("Delivery", f"{vendor.get('delivery_days','')} days ARO" if vendor.get("delivery_days") else ""),
        ("Prepared By", vendor.get("prepared_by","")),
    ]:
        if value:
            mp=rc.add_paragraph()
            mp.paragraph_format.left_indent=Pt(8)
            run(mp,f"{label}: ",bold=True,size=9,color=RGBColor(0x33,0x33,0x33))
            run(mp,value,size=9,color=RGBColor(0x33,0x33,0x33))
    rc.add_paragraph()
    doc.add_paragraph()
    # Solicitation info
    heading("SOLICITATION INFORMATION")
    fields=[
        ("Project / Subject",    solicitation.get("project_title")),
        ("Solicitation Number",  solicitation.get("solicitation_number")),
        ("Solicitation Type",    solicitation.get("solicitation_type")),
        ("Issuing Agency",       solicitation.get("issuing_agency")),
        ("Contracting Office",   solicitation.get("contracting_office_address")),
        ("Response Due Date",    solicitation.get("due_date")),
        ("NAICS Code",           solicitation.get("naics_code")),
        ("PSC Code",             solicitation.get("psc_code")),
        ("Set-Aside",            solicitation.get("set_aside")),
        ("Place of Performance", solicitation.get("place_of_performance")),
        ("Period of Performance",solicitation.get("period_of_performance")),
    ]
    fields=[(l,v) for l,v in fields if v]
    if fields:
        st=doc.add_table(rows=len(fields),cols=2); st.style="Table Grid"; st.autofit=False
        st.columns[0].width=Inches(2.0); st.columns[1].width=Inches(4.0)
        st.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i,(label,value) in enumerate(fields):
            row_keep(st.rows[i])
            lc2=st.cell(i,0); rc2=st.cell(i,1)
            bg(lc2,"F0F0F0")
            lp2=lc2.paragraphs[0]; lp2.paragraph_format.left_indent=Pt(4)
            run(lp2,label,bold=True,size=9,color=NAVY)
            rp2=rc2.paragraphs[0]; rp2.paragraph_format.left_indent=Pt(4)
            run(rp2,str(value)[:200],size=9,color=DGRAY)

    # Scope
    scope=solicitation.get("scope_of_work","")
    if scope:
        heading("SCOPE OF WORK / REQUIREMENT")
        sp=doc.add_paragraph(); sp.paragraph_format.left_indent=Inches(0.1)
        run(sp, scope[:2000], size=9.5, color=DGRAY)

    # Line items
    heading("QUOTE DETAILS")
    cw=[Inches(0.35),Inches(2.1),Inches(0.55),Inches(0.6),Inches(0.55),Inches(0.85),Inches(1.0)]
    hdrs=["#","Description / Item","Size/Type","UOM","Qty","Unit Price","Total"]
    lt=doc.add_table(rows=1+len(line_items)+1,cols=7); lt.style="Table Grid"; lt.autofit=False
    lt.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci,w in enumerate(cw): lt.columns[ci].width=w
    hr=lt.rows[0]
    row_keep(hr, header=True)
    for ci,h in enumerate(hdrs):
        c=hr.cells[ci]; bg(c,"1A1A1A"); c.width=cw[ci]
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run(p,h,bold=True,size=9,color=WHITE)
    grand=0.0; has_any_price=False
    AL=[WD_ALIGN_PARAGRAPH.CENTER,WD_ALIGN_PARAGRAPH.LEFT,WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,WD_ALIGN_PARAGRAPH.CENTER,WD_ALIGN_PARAGRAPH.RIGHT,
        WD_ALIGN_PARAGRAPH.RIGHT]
    for i,item in enumerate(line_items):
        row=lt.rows[i+1]; row_keep(row)
        qty_n = fmt_num(item.get("qty"))
        up_n  = fmt_num(item.get("unit_price"))
        total_n = (qty_n * up_n) if (qty_n is not None and up_n is not None) else None
        if total_n is not None: grand += total_n; has_any_price = True
        bcolor="FFFFFF" if i%2==0 else "F7F9FC"
        qty_s   = (str(int(qty_n)) if qty_n is not None and qty_n == int(qty_n) else str(qty_n)) if qty_n is not None else "N/A"
        up_s    = f"${up_n:,.2f}" if up_n is not None else "N/A"
        total_s = f"${total_n:,.2f}" if total_n is not None else "N/A"
        desc    = item.get("description","") or "N/A"
        size    = item.get("size","") or "N/A"
        unit    = item.get("unit","EA") or "EA"
        vals=[str(i+1), desc, size, unit, qty_s, up_s, total_s]
        for ci,(val,al,w) in enumerate(zip(vals,AL,cw)):
            c=row.cells[ci]; bg(c,bcolor); c.width=w
            p=c.paragraphs[0]; p.alignment=al; p.paragraph_format.left_indent=Pt(3)
            run(p,val,size=9,color=DGRAY)
    tr=lt.rows[-1]; row_keep(tr)
    for ci in range(7): bg(tr.cells[ci],"F0F0F0"); tr.cells[ci].width=cw[ci]
    tr.cells[0].merge(tr.cells[4])
    tp=tr.cells[0].paragraphs[0]; tp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    run(tp,"GRAND TOTAL",bold=True,size=10,color=NAVY)
    freight=float(vendor.get("freight",0) or 0)
    tax_rate=float(vendor.get("tax_rate",0) or 0)
    tax=grand*(tax_rate/100); final=grand+freight+tax
    gp=tr.cells[-1].paragraphs[0]; gp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    final_s = f"${final:,.2f}" if (has_any_price or freight>0) else "N/A"
    run(gp,final_s,bold=True,size=11,color=NAVY)

    # Notes & Terms
    notes=vendor.get("notes",""); terms=vendor.get("terms","")
    if notes or terms:
        heading("NOTES & TERMS")
        if notes:
            np=doc.add_paragraph(); np.paragraph_format.left_indent=Inches(0.1)
            run(np,"Notes: ",bold=True,size=9.5,color=NAVY); run(np,notes,size=9.5,color=DGRAY)
        if terms:
            tp2=doc.add_paragraph(); tp2.paragraph_format.left_indent=Inches(0.1)
            run(tp2,"Terms: ",bold=True,size=9.5,color=NAVY); run(tp2,terms,size=9.5,color=DGRAY)

    # Option years summary
    option_years = vendor.get("option_years", [])
    if vendor.get("option_years_enabled") and option_years and has_any_price:
        base_total = grand + freight + tax
        heading("OPTION YEAR PRICING SUMMARY")
        oy_rows = [("Base Year (Year 1)", base_total)]
        for oy in option_years:
            pct = float(oy.get("pct", 0) or 0)
            oy_rows.append((oy.get("label","Option Year"), base_total * (1 + pct / 100)))
        oy_rows.append(("TOTAL — All Periods", sum(r[1] for r in oy_rows)))
        ot = doc.add_table(rows=len(oy_rows), cols=2); ot.style="Table Grid"; ot.autofit=False
        ot.columns[0].width = Inches(3.5); ot.columns[1].width = Inches(2.5)
        ot.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i,(period,total_oy) in enumerate(oy_rows):
            row_keep(ot.rows[i])
            is_last = (i == len(oy_rows) - 1)
            lc3=ot.cell(i,0); rc3=ot.cell(i,1)
            bg(lc3,"F0F0F0" if (i==0 or is_last) else "FFFFFF")
            bg(rc3,"F0F0F0" if (i==0 or is_last) else "FFFFFF")
            lp3=lc3.paragraphs[0]; lp3.paragraph_format.left_indent=Pt(4)
            run(lp3,period,bold=is_last,size=9,color=NAVY)
            rp3=rc3.paragraphs[0]; rp3.alignment=WD_ALIGN_PARAGRAPH.RIGHT
            run(rp3,f"${total_oy:,.2f}",bold=is_last,size=9,color=NAVY if is_last else DGRAY)

    # Signature
    heading("AUTHORIZED SIGNATURE", sb=18)
    sigt=doc.add_table(rows=1,cols=2); sigt.style="Table Grid"; sigt.autofit=False
    sigt.columns[0].width=Inches(3.0); sigt.columns[1].width=Inches(3.0)
    sigt.alignment = WD_TABLE_ALIGNMENT.CENTER
    row_keep(sigt.rows[0])
    lsc=sigt.cell(0,0); rsc=sigt.cell(0,1)
    def sigline(cell,label,value=""):
        p=cell.add_paragraph(); p.paragraph_format.left_indent=Pt(6); p.paragraph_format.space_after=Pt(10)
        run(p,f"{label}: ",bold=True,size=9,color=NAVY); run(p,value or "_"*28,size=9,color=DGRAY)
    sigline(lsc,"Authorized Signature"); sigline(lsc,"Printed Name",vendor.get("prepared_by",""))
    sigline(lsc,"Title",vendor.get("title","")); sigline(lsc,"Date",today)
    sigline(rsc,"Company",vendor.get("company_name","")); sigline(rsc,"Phone",vendor.get("phone",""))
    sigline(rsc,"Email",vendor.get("email","")); sigline(rsc,"SAM UEI",vendor.get("sam_uei",""))
    doc.add_paragraph()
    fp=doc.add_paragraph(); fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run(fp,f"Quote valid for {vendor.get('validity_period','30 days')} from {today}.",
        size=8,color=RGBColor(0x99,0x99,0x99),italic=True)

    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf.read()

# ── ROUTES ───────────────────────────────────────────────────────────────────

@app.route("/parse", methods=["POST","OPTIONS"])
def parse_route():
    if request.method=="OPTIONS": return jsonify({}),200
    try:
        if "file" not in request.files:
            return jsonify({"error":"No file uploaded"}),400
        file=request.files["file"]; api_key=request.form.get("api_key","")
        suffix=Path(file.filename).suffix
        with tempfile.NamedTemporaryFile(delete=False,suffix=suffix) as tmp:
            file.save(tmp.name); tmp_path=tmp.name
        try:
            text=parse_document(tmp_path)
            if not text.strip():
                return jsonify({"error":"Could not extract text from document."}),400
            data=extract_data(text,api_key)
            return jsonify({"success":True,"data":data})
        finally:
            os.unlink(tmp_path)
    except Exception as e:
        traceback.print_exc(); return jsonify({"error":str(e)}),500

@app.route("/generate_quote", methods=["POST","OPTIONS"])
def gen_route():
    if request.method=="OPTIONS": return jsonify({}),200
    try:
        body=request.get_json()
        docx_bytes=generate_quote(body.get("solicitation",{}),body.get("vendor",{}),body.get("line_items",[]))
        sol=body.get("solicitation",{}).get("solicitation_number","quote").replace("/","-")
        return send_file(io.BytesIO(docx_bytes),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True, download_name=f"Quote_{sol}.docx")
    except Exception as e:
        traceback.print_exc(); return jsonify({"error":str(e)}),500

@app.route("/sam_lookup", methods=["POST","OPTIONS"])
def sam_lookup():
    if request.method=="OPTIONS": return jsonify({}),200
    try:
        body=request.get_json()
        notice_id=body.get("notice_id","").strip()
        api_key=body.get("api_key","").strip()
        if not notice_id: return jsonify({"error":"Notice ID is required"}),400
        if not api_key:   return jsonify({"error":"SAM.gov API key is required"}),400

        from datetime import datetime,timedelta
        import urllib.request as _ur, urllib.parse as _up
        today=datetime.now(); ago=today-timedelta(days=365)
        params=_up.urlencode({"api_key":api_key,"noticeid":notice_id,
                               "postedFrom":ago.strftime("%m/%d/%Y"),
                               "postedTo":today.strftime("%m/%d/%Y"),"limit":1})
        url=f"https://api.sam.gov/opportunities/v2/search?{params}"
        req=_ur.Request(url,headers={"Accept":"application/json","User-Agent":"SolicitationQuoter/1.0"})
        with _ur.urlopen(req,timeout=15) as resp:
            data=json.loads(resp.read())

        opps=data.get("opportunitiesData",[])
        if not opps: return jsonify({"error":f"No opportunity found for: {notice_id}"}),404
        opp=opps[0]

        desc_raw=opp.get("description","")
        desc=re.sub(r"\s+"," ",re.sub(r"<[^>]+>"," ",desc_raw)).strip()

        solicitation={
            "solicitation_number": opp.get("solicitationNumber") or opp.get("noticeId",""),
            "project_title":       opp.get("title",""),
            "solicitation_type":   opp.get("type",""),
            "issuing_agency":      opp.get("fullParentPathName",""),
            "due_date":            opp.get("responseDeadLine",""),
            "posting_date":        opp.get("postedDate",""),
            "naics_code":          opp.get("naicsCode",""),
            "psc_code":            opp.get("classificationCode",""),
            "set_aside":           opp.get("typeOfSetAsideDescription") or opp.get("typeOfSetAside",""),
            "scope_of_work":       desc[:3000],
        }
        pop=opp.get("placeOfPerformance",{})
        if pop:
            city=pop.get("city",{}).get("name",""); state=pop.get("state",{}).get("code","")
            if city or state: solicitation["place_of_performance"]=f"{city}, {state}".strip(", ")
        contacts=opp.get("pointOfContact",[])
        if contacts:
            c=contacts[0]
            solicitation["contact_name"]=c.get("fullName","")
            solicitation["contact_email"]=c.get("email","")
            solicitation["contact_phone"]=c.get("phone","")
        office=opp.get("officeAddress",{})
        if office:
            solicitation["contracting_office_address"]=" ".join(
                filter(None,[office.get("city",""),office.get("state",""),office.get("zipcode","")]))

        solicitation["_method"]="sam_gov"
        # Also run rule-based extraction on desc for quantities
        rules=extract(desc)
        for k,v in rules.items():
            if v and k not in solicitation: solicitation[k]=v

        line_items=extract_line_items(solicitation, desc)
        return jsonify({"success":True,"data":solicitation,"line_items":line_items})

    except Exception as e:
        import urllib.error
        if isinstance(e, urllib.error.HTTPError):
            if e.code == 403:
                return jsonify({"error": "Invalid or expired SAM.gov API key."}), 400
            return jsonify({"error": f"SAM.gov returned an error (HTTP {e.code}). Try again later."}), 502
        if isinstance(e, urllib.error.URLError):
            if 'timed out' in str(e).lower() or isinstance(getattr(e, 'reason', None), Exception) and 'timed out' in str(e.reason).lower():
                return jsonify({"error": "SAM.gov did not respond in time. Check your connection and try again."}), 504
            return jsonify({"error": f"Could not reach SAM.gov. Check your internet connection."}), 503
        if isinstance(e, (ValueError, KeyError)):
            return jsonify({"error": "SAM.gov returned an unexpected response. Try again later."}), 502
        traceback.print_exc(); return jsonify({"error": str(e)}), 500

if __name__=="__main__":
    port=int(os.environ.get("PORT",5199))
    print(f"[SolicitationQuoter] Running on http://localhost:{port}")
    app.run(host="127.0.0.1",port=port,debug=False)
